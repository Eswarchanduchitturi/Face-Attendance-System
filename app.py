from flask import Flask, render_template, Response, request, redirect, url_for, send_file, jsonify, session
from werkzeug.utils import secure_filename

import cv2
import calendar
import sqlite3
import datetime
import os
import io
import math
import threading
import uuid
from collections import deque, Counter

import numpy as np
from PIL import Image, ImageDraw, ImageFont
import time
from routes.training import create_training_blueprint
from face_utils import (
    load_cascade_classifiers,
    detect_faces_optimized,
    extract_and_prepare_face,
    preprocess_face,
    resize_face_with_padding,
    predict_face_with_flip,
    verify_prediction_with_gallery,
    face_similarity,
    draw_face_rectangle,
    is_blurry,
    FACE_WIDTH,
    FACE_HEIGHT,
    CONFIDENCE_THRESHOLD,
    MIN_STABLE_FRAMES
)



# System metrics
camera_status = "OFF"
fps_value = 0
frame_count = 0
start_time = time.time()

# Model performance metrics
success_recognition = 0
failed_recognition = 0
confidence_scores = []




db = sqlite3.connect("database/attendance.db")

db.execute("INSERT OR IGNORE INTO users VALUES (1,'Eswar','admin')")
db.execute("INSERT OR IGNORE INTO users VALUES (2,'Ravi','employee')")
db.execute("""
    UPDATE users
    SET role='employee'
    WHERE COALESCE(LOWER(role), '') IN ('student', 'faculty', 'user', '')
""")

db.commit()
db.close()

# Load cascade classifiers using improved function
face_cascade, eye_cascade = load_cascade_classifiers()

if face_cascade is None:
    print("⚠️ WARNING: Face cascade classifier not available. Face detection will not work.")

blink_detected = False
head_moved = False
prev_face_x = None
camera = None
camera_enabled = False
attendance_status = ""
attendance_mode = "checkin"   # default mode
last_attendance_time = 0
ATTENDANCE_COOLDOWN = 3  # seconds






from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch
from flask import session
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor



# ------------------ Flask App ------------------
app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY') or 'dev-secret-key-change-in-production'

# ------------------ Config ------------------
CAMERA_INDEX = 0
THRESHOLD = CONFIDENCE_THRESHOLD  # Use imported constant
TOTAL_EMPLOYEES = 10    # change later (will auto-calc with users table)
RECOGNIZE_COOLDOWN_FRAMES = 5  # Wait N frames before re-triggering recognition
TARGET_STREAM_FPS = 15
SMALL_FRAME_SCALE = 0.6  # Downscale for faster detection
REQUIRED_STABLE_MATCHES = max(4, MIN_STABLE_FRAMES)
MIN_IMAGES_PER_USER = 12
MIN_FULL_DAY_HOURS = 7.0
DUPLICATE_MIN_IMAGES_PER_USER = MIN_IMAGES_PER_USER


def get_default_base_salary_for_role(role):
    role_normalized = (role or "employee").strip().lower()
    return 50000.0 if role_normalized == "admin" else 40000.0

# Enrollment duplicate-face guard rails
DUPLICATE_CONF_THRESHOLD = min(45, CONFIDENCE_THRESHOLD)
STRONG_DUPLICATE_CONF = min(40, CONFIDENCE_THRESHOLD)
DUPLICATE_MIN_HITS = 5
DUPLICATE_REQUIRED_RATIO = 0.55
DUPLICATE_MIN_CONF_MARGIN = 4.0

# Mask-aware recognition defaults (used by admin config page)
MASK_AWARE_ENABLED = False
MASK_THRESHOLD_BOOST = 8
MASK_REQUIRED_STABLE_MATCHES = 4
MASK_MIN_UPPER_CORR = 0.45
MASK_MIN_UPPER_GAP = 0.03
MASK_MAX_UPPER_LBP_DIST = 1.30
MASK_MIN_UPPER_LBP_MARGIN = 0.08

# Optional iris verification defaults
IRIS_VERIFY_ENABLED = False
IRIS_MIN_SCORE = 0.12
IRIS_MIN_GAP = 0.03
IRIS_MIN_VOTES = 1

face_samples_cache = {}
face_samples_cache_model_mtime = None
recognizer_model_mtime = None
user_cache = {}
user_cache_loaded_at = 0.0
training_image_count_cache = {}
training_image_count_loaded_at = 0.0
MODEL_PATH = "TrainingImageLabel/Trainner.yml"

# Load face cascade classifier using improved function
if face_cascade is None:
    print("❌ ERROR: Could not load face cascade classifier!")
else:
    print("✅ Face cascade classifier ready for detection")

# Load LBPH Face Recognizer and trained model
recognizer = cv2.face.LBPHFaceRecognizer_create()
if os.path.exists(MODEL_PATH):
    try:
        recognizer.read(MODEL_PATH)
        recognizer_model_mtime = os.path.getmtime(MODEL_PATH)
        print(f"✅ Trained model loaded successfully")
    except Exception as e:
        print(f"⚠️ Could not load trained model: {e}")
        print("   Please enroll users and train the model first.")
else:
    print("⚠️ No trained model found at TrainingImageLabel/Trainner.yml")
    print("   Please enroll users and train the model first.")

# ------------------ Database ------------------
def get_db():
    return sqlite3.connect("database/attendance.db")


def ensure_payroll_schema():
    db = get_db()
    cursor = db.cursor()

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS payroll_settings (
        user_id INTEGER PRIMARY KEY,
        monthly_salary REAL DEFAULT 0,
        hourly_rate REAL DEFAULT 0,
        workdays_per_month INTEGER DEFAULT 0,
        standard_hours REAL DEFAULT 8,
        overtime_multiplier REAL DEFAULT 1.5,
        deduction_per_absent_day REAL DEFAULT 0,
        bonus_amount REAL DEFAULT 0,
        updated_at TEXT,
        FOREIGN KEY (user_id) REFERENCES users(id)
    )
    """)

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS payroll_records (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        month TEXT NOT NULL,
        year INTEGER NOT NULL,
        present_days INTEGER DEFAULT 0,
        absent_days INTEGER DEFAULT 0,
        expected_workdays INTEGER DEFAULT 0,
        total_hours REAL DEFAULT 0,
        overtime_hours REAL DEFAULT 0,
        base_salary REAL DEFAULT 0,
        overtime_pay REAL DEFAULT 0,
        deductions REAL DEFAULT 0,
        bonuses REAL DEFAULT 0,
        gross_salary REAL DEFAULT 0,
        net_salary REAL DEFAULT 0,
        generated_at TEXT,
        status TEXT DEFAULT 'generated',
        UNIQUE(user_id, month),
        FOREIGN KEY (user_id) REFERENCES users(id)
    )
    """)

    db.commit()
    db.close()


def ensure_leave_schema():
    db = get_db()
    cursor = db.cursor()

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS leave_requests (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        emp_id TEXT NOT NULL,
        subject TEXT NOT NULL,
        leave_type TEXT DEFAULT 'General',
        start_date TEXT NOT NULL,
        end_date TEXT NOT NULL,
        total_days INTEGER NOT NULL,
        reason TEXT NOT NULL,
        request_status TEXT DEFAULT 'Pending',
        admin_comment TEXT,
        applied_at TEXT,
        reviewed_at TEXT,
        reviewed_by TEXT,
        FOREIGN KEY (emp_id) REFERENCES users(id)
    )
    """)

    db.commit()
    db.close()


def get_month_key(target_date=None):
    if target_date is None:
        target_date = datetime.date.today()
    return target_date.strftime("%Y-%m")


def get_month_bounds(month_key):
    year, month = (int(value) for value in month_key.split("-"))
    last_day = calendar.monthrange(year, month)[1]
    start_date = datetime.date(year, month, 1)
    end_date = datetime.date(year, month, last_day)
    return year, month, start_date, end_date


def count_workdays_in_month(year, month):
    first_day = datetime.date(year, month, 1)
    last_day = datetime.date(year, month, calendar.monthrange(year, month)[1])
    workdays = 0
    current_day = first_day

    while current_day <= last_day:
        if current_day.weekday() < 5:
            workdays += 1
        current_day += datetime.timedelta(days=1)

    return workdays


def count_completed_workdays_in_month(year, month, reference_date=None):
    """Count workdays completed so far for a given month (Mon-Fri only)."""
    if reference_date is None:
        reference_date = datetime.date.today()

    first_day = datetime.date(year, month, 1)
    last_day = datetime.date(year, month, calendar.monthrange(year, month)[1])

    if reference_date < first_day:
        return 0

    effective_end = min(reference_date, last_day)
    workdays = 0
    current_day = first_day

    while current_day <= effective_end:
        if current_day.weekday() < 5:
            workdays += 1
        current_day += datetime.timedelta(days=1)

    return workdays


def count_workdays_in_range(start_date, end_date):
    if start_date > end_date:
        start_date, end_date = end_date, start_date

    workdays = 0
    current_day = start_date

    while current_day <= end_date:
        if current_day.weekday() < 5:
            workdays += 1
        current_day += datetime.timedelta(days=1)

    return workdays


def get_default_report_dates():
    today = datetime.date.today()
    start_date = today.replace(day=1)
    return start_date.isoformat(), today.isoformat()


def parse_int_or_none(value):
    if value is None:
        return None

    raw_value = str(value).strip()
    if not raw_value:
        return None

    try:
        return int(raw_value)
    except ValueError:
        return None


def normalize_report_format(raw_format):
    value = (raw_format or "pdf").strip().lower()
    if value in {"xls", "excel"}:
        return "xlsx"
    if value in {"doc", "word"}:
        return "docx"
    if value in {"png", "jpg", "jpeg", "image"}:
        return "png"
    return value if value in {"pdf", "xlsx", "docx", "png"} else "pdf"


def get_attendance_report_scope_label(employee_id, employees_map):
    if employee_id is None:
        return "All Employees"

    employee = employees_map.get(employee_id)
    if not employee:
        return f"Employee #{employee_id}"

    return f"{employee['name']} (ID {employee_id})"


def build_custom_attendance_report_data(db, employee_id=None, start_date=None, end_date=None):
    employees = {
        row[0]: {"name": row[1], "role": row[2] or "employee"}
        for row in db.execute("SELECT id, name, role FROM users ORDER BY name").fetchall()
    }

    params = [start_date.isoformat(), end_date.isoformat()]
    employee_filter_sql = ""
    if employee_id is not None:
        employee_filter_sql = " AND CAST(attendance.emp_id AS INTEGER) = ? "
        params.append(employee_id)

    rows = db.execute(f"""
        SELECT attendance.id,
               attendance.emp_id,
               COALESCE(users.name, 'Unknown') AS name,
               COALESCE(users.role, 'employee') AS role,
               attendance.date,
               attendance.day,
               attendance.checkin_time,
               attendance.checkout_time,
               ROUND(COALESCE(attendance.worked_hours, 0), 2) AS worked_hours,
               COALESCE(attendance.status, 'Present') AS status,
               attendance.checkin_image,
               attendance.checkout_image
        FROM attendance
        LEFT JOIN users ON CAST(attendance.emp_id AS INTEGER) = users.id
        WHERE attendance.date BETWEEN ? AND ?
        {employee_filter_sql}
        ORDER BY attendance.date ASC, attendance.checkin_time ASC, attendance.id ASC
    """, tuple(params)).fetchall()

    employees_in_scope = {}
    detail_rows = []

    for row in rows:
        emp_id = int(row[1]) if row[1] is not None else None
        if emp_id is None:
            continue

        employee_info = employees.get(emp_id, {"name": row[2], "role": row[3]})
        if emp_id not in employees_in_scope:
            employees_in_scope[emp_id] = {
                "user_id": emp_id,
                "name": employee_info["name"],
                "role": employee_info["role"],
                "daily_hours": {},
                "sessions": 0,
                "present_days": set(),
            }

        detail_rows.append({
            "user_id": emp_id,
            "name": employee_info["name"],
            "role": employee_info["role"],
            "date": row[4],
            "day": row[5],
            "checkin_time": row[6] or "-",
            "checkout_time": row[7] or "-",
            "worked_hours": float(row[8] or 0),
            "status": row[9] or "Present",
        })

        if row[4]:
            daily_hours = employees_in_scope[emp_id]["daily_hours"]
            daily_hours[row[4]] = daily_hours.get(row[4], 0.0) + float(row[8] or 0)
            employees_in_scope[emp_id]["sessions"] += 1
            if float(row[8] or 0) > 0:
                employees_in_scope[emp_id]["present_days"].add(row[4])

    summary_rows = []
    for emp_id in sorted(employees_in_scope.keys()):
        employee = employees_in_scope[emp_id]
        summary = calculate_attendance_salary_projection(
            db=db,
            employee_id=emp_id,
            employee_name=employee["name"],
            role=employee["role"],
            daily_hours=employee["daily_hours"],
            start_date=start_date,
            end_date=end_date,
        )
        summary_rows.append(summary)

    if employee_id is not None and employee_id not in employees and not detail_rows:
        return {
            "employees": employees,
            "detail_rows": [],
            "summary_rows": [],
            "scope_label": get_attendance_report_scope_label(employee_id, employees),
            "start_date": start_date,
            "end_date": end_date,
            "selected_employee_id": employee_id,
            "selected_employee_name": None,
        }

    selected_employee_name = None
    if employee_id is not None:
        selected_employee_name = employees.get(employee_id, {}).get("name")

    return {
        "employees": employees,
        "detail_rows": detail_rows,
        "summary_rows": summary_rows,
        "scope_label": get_attendance_report_scope_label(employee_id, employees),
        "start_date": start_date,
        "end_date": end_date,
        "selected_employee_id": employee_id,
        "selected_employee_name": selected_employee_name,
    }


def calculate_attendance_salary_projection(db, employee_id, employee_name, role, daily_hours, start_date, end_date):
    settings = get_payroll_settings(db, employee_id, get_month_key(start_date))
    standard_hours = float(settings["standard_hours"] or 8)
    expected_workdays = settings["workdays_per_month"] or count_workdays_in_range(start_date, end_date)
    present_days = len([day for day, hours in daily_hours.items() if hours > 0])
    total_hours = round(sum(float(hours or 0) for hours in daily_hours.values()), 2)

    paid_regular_hours = 0.0
    overtime_hours = 0.0
    paid_days = 0
    for day_hours in daily_hours.values():
        day_hours = float(day_hours or 0)
        if day_hours >= MIN_FULL_DAY_HOURS:
            paid_regular_hours += standard_hours
            paid_days += 1
        else:
            paid_regular_hours += max(day_hours, 0.0)
        overtime_hours += max(day_hours - standard_hours, 0.0)

    default_base_salary = get_default_base_salary_for_role(role)
    monthly_salary = max(float(settings["monthly_salary"] or 0), default_base_salary)
    hourly_rate = float(settings["hourly_rate"] or 0)
    overtime_multiplier = float(settings["overtime_multiplier"] or 1.5)
    bonus_amount = float(settings["bonus_amount"] or 0)
    deduction_per_absent_day = float(settings["deduction_per_absent_day"] or 0)

    if monthly_salary > 0 and hourly_rate <= 0:
        effective_hourly_rate = monthly_salary / max(expected_workdays * standard_hours, 1)
    else:
        effective_hourly_rate = max(hourly_rate, 0)

    absent_days = max(expected_workdays - present_days, 0)
    base_salary = round(max(paid_regular_hours, 0.0) * effective_hourly_rate, 2)
    overtime_pay = round(max(overtime_hours, 0.0) * effective_hourly_rate * overtime_multiplier, 2)
    deductions = round(absent_days * deduction_per_absent_day, 2)
    gross_salary = round(base_salary + overtime_pay + bonus_amount, 2)
    net_salary = round(max(gross_salary - deductions, 0.0), 2)

    return {
        "user_id": employee_id,
        "name": employee_name,
        "role": role,
        "expected_workdays": expected_workdays,
        "paid_days": paid_days,
        "present_days": present_days,
        "absent_days": absent_days,
        "total_hours": total_hours,
        "overtime_hours": round(overtime_hours, 2),
        "base_salary": round(base_salary, 2),
        "overtime_pay": round(overtime_pay, 2),
        "deductions": round(deductions, 2),
        "gross_salary": round(gross_salary, 2),
        "net_salary": round(net_salary, 2),
        "standard_hours": standard_hours,
        "effective_hourly_rate": round(effective_hourly_rate, 2),
        "bonus_amount": round(bonus_amount, 2),
        "daily_hours": daily_hours,
    }


def build_report_download_name(scope_label, start_date, end_date, extension):
    safe_scope = secure_filename(scope_label) or "attendance_report"
    return f"{safe_scope}_{start_date.isoformat()}_to_{end_date.isoformat()}.{extension}"


def render_attendance_report_pdf(report_data):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=0.45 * inch,
        bottomMargin=0.45 * inch,
        leftMargin=0.45 * inch,
        rightMargin=0.45 * inch,
    )
    story = []
    styles = getSampleStyleSheet()

    title_style = styles["Title"]
    title_style.textColor = colors.HexColor("#0f1720")

    subtitle_style = styles["Normal"]
    subtitle_style.textColor = colors.HexColor("#5b6b79")

    story.append(Paragraph("Attendance & Salary Report", title_style))
    story.append(Paragraph(f"Scope: <b>{report_data['scope_label']}</b>", subtitle_style))
    story.append(Paragraph(
        f"Period: <b>{report_data['start_date'].isoformat()}</b> to <b>{report_data['end_date'].isoformat()}</b>",
        subtitle_style,
    ))
    story.append(Paragraph(
        f"Generated on: {datetime.datetime.now().strftime('%B %d, %Y at %I:%M %p')}",
        subtitle_style,
    ))
    story.append(Spacer(1, 0.18 * inch))

    if report_data["summary_rows"]:
        summary_header = [[
            Paragraph("<b>Employee</b>", styles["BodyText"]),
            Paragraph("<b>Role</b>", styles["BodyText"]),
            Paragraph("<b>Present</b>", styles["BodyText"]),
            Paragraph("<b>Hours</b>", styles["BodyText"]),
            Paragraph("<b>Salary</b>", styles["BodyText"]),
        ]]
        for row in report_data["summary_rows"]:
            summary_header.append([
                f"{row['name']} ({row['user_id']})",
                row["role"],
                str(row["present_days"]),
                f"{row['total_hours']:.2f}",
                f"{row['net_salary']:.2f}",
            ])

        summary_table = Table(summary_header, colWidths=[2.4 * inch, 1.0 * inch, 0.8 * inch, 0.9 * inch, 1.1 * inch])
        summary_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0d6efd")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#f5f9ff")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#ffffff"), colors.HexColor("#eef5ff")]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d7e2f0")),
            ("TOPPADDING", (0, 0), (-1, 0), 10),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 10),
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 0.25 * inch))

        if len(report_data["summary_rows"]) == 1:
            summary = report_data["summary_rows"][0]
            detail_card_data = [[
                Paragraph(f"<b>Employee</b><br/>{summary['name']} ({summary['user_id']})", styles["BodyText"]),
                Paragraph(f"<b>Present Days</b><br/>{summary['present_days']}", styles["BodyText"]),
                Paragraph(f"<b>Total Hours</b><br/>{summary['total_hours']:.2f}", styles["BodyText"]),
                Paragraph(f"<b>Net Salary</b><br/>{summary['net_salary']:.2f}", styles["BodyText"]),
            ]]
            detail_card_table = Table(detail_card_data, colWidths=[2.3 * inch, 1.3 * inch, 1.3 * inch, 1.3 * inch])
            detail_card_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f8fbff")),
                ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#d9e8fb")),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d9e8fb")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("TOPPADDING", (0, 0), (-1, -1), 12),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
            ]))
            story.append(detail_card_table)
            story.append(Spacer(1, 0.2 * inch))

    if report_data["detail_rows"]:
        story.append(Paragraph("<b>Attendance Sessions</b>", styles["Heading2"]))
        story.append(Spacer(1, 0.08 * inch))

        detail_data = [["User ID", "Name", "Date", "Day", "Check-In", "Check-Out", "Hours", "Status"]]
        for row in report_data["detail_rows"]:
            detail_data.append([
                str(row["user_id"]),
                row["name"],
                row["date"],
                row["day"] or "-",
                row["checkin_time"],
                row["checkout_time"],
                f"{row['worked_hours']:.2f}",
                row["status"],
            ])

        detail_table = Table(detail_data, colWidths=[0.65 * inch, 1.2 * inch, 0.95 * inch, 0.75 * inch, 0.95 * inch, 0.95 * inch, 0.65 * inch, 0.95 * inch])
        detail_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#198754")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 8.5),
            ("FONTSIZE", (0, 1), (-1, -1), 8),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5fbf7")]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d4e7dc")),
            ("TOPPADDING", (0, 0), (-1, 0), 9),
            ("BOTTOMPADDING", (0, 0), (-1, 0), 9),
        ]))
        story.append(detail_table)
    else:
        story.append(Paragraph("No attendance sessions found for the selected employee and date range.", styles["BodyText"]))

    if report_data["summary_rows"]:
        salary_note = ", ".join(
            f"{row['name']}: {row['net_salary']:.2f}" for row in report_data["summary_rows"][:3]
        )
        story.append(Spacer(1, 0.16 * inch))
        story.append(Paragraph(f"<i>Salary preview: {salary_note}</i>", styles["BodyText"]))

    doc.build(story)
    buffer.seek(0)
    return buffer


def render_attendance_report_xlsx(report_data):
    workbook = Workbook()
    summary_sheet = workbook.active
    summary_sheet.title = "Summary"
    details_sheet = workbook.create_sheet("Attendance")

    header_fill = PatternFill("solid", fgColor="0D6EFD")
    green_fill = PatternFill("solid", fgColor="198754")
    light_fill = PatternFill("solid", fgColor="F5FBFF")
    alt_fill = PatternFill("solid", fgColor="EEF5FF")
    header_font = Font(color="FFFFFF", bold=True)
    body_font = Font(color="111111")
    center = Alignment(horizontal="center", vertical="center")
    left = Alignment(horizontal="left", vertical="center")
    thin_border = Border(
        left=Side(style="thin", color="D7E2F0"),
        right=Side(style="thin", color="D7E2F0"),
        top=Side(style="thin", color="D7E2F0"),
        bottom=Side(style="thin", color="D7E2F0"),
    )

    summary_sheet["A1"] = "Attendance & Salary Report"
    summary_sheet["A1"].font = Font(bold=True, size=16, color="0F1720")
    summary_sheet.merge_cells("A1:F1")
    summary_sheet["A2"] = f"Scope: {report_data['scope_label']}"
    summary_sheet["A3"] = f"Period: {report_data['start_date'].isoformat()} to {report_data['end_date'].isoformat()}"
    summary_sheet["A4"] = f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"

    summary_start_row = 6
    summary_headers = ["User ID", "Name", "Role", "Present Days", "Hours", "Estimated Net Salary"]
    for col_index, value in enumerate(summary_headers, start=1):
        cell = summary_sheet.cell(row=summary_start_row, column=col_index, value=value)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center
        cell.border = thin_border

    for row_index, row in enumerate(report_data["summary_rows"], start=summary_start_row + 1):
        values = [row["user_id"], row["name"], row["role"], row["present_days"], round(row["total_hours"], 2), round(row["net_salary"], 2)]
        for col_index, value in enumerate(values, start=1):
            cell = summary_sheet.cell(row=row_index, column=col_index, value=value)
            cell.font = body_font
            cell.alignment = left if col_index in {2, 3} else center
            cell.border = thin_border
            cell.fill = light_fill if row_index % 2 else alt_fill

    if report_data["summary_rows"]:
        summary_total_row = summary_start_row + len(report_data["summary_rows"]) + 2
        summary_sheet.cell(row=summary_total_row, column=1, value="Totals").font = Font(bold=True)
        summary_sheet.cell(row=summary_total_row, column=4, value=sum(row["present_days"] for row in report_data["summary_rows"]))
        summary_sheet.cell(row=summary_total_row, column=5, value=round(sum(row["total_hours"] for row in report_data["summary_rows"]), 2))
        summary_sheet.cell(row=summary_total_row, column=6, value=round(sum(row["net_salary"] for row in report_data["summary_rows"]), 2))

    details_headers = ["User ID", "Name", "Role", "Date", "Day", "Check-In", "Check-Out", "Hours", "Status"]
    for col_index, value in enumerate(details_headers, start=1):
        cell = details_sheet.cell(row=1, column=col_index, value=value)
        cell.fill = green_fill
        cell.font = header_font
        cell.alignment = center
        cell.border = thin_border

    for row_index, row in enumerate(report_data["detail_rows"], start=2):
        values = [
            row["user_id"],
            row["name"],
            row["role"],
            row["date"],
            row["day"],
            row["checkin_time"],
            row["checkout_time"],
            round(row["worked_hours"], 2),
            row["status"],
        ]
        for col_index, value in enumerate(values, start=1):
            cell = details_sheet.cell(row=row_index, column=col_index, value=value)
            cell.font = body_font
            cell.alignment = left if col_index in {2, 3, 9} else center
            cell.border = thin_border
            cell.fill = light_fill if row_index % 2 else alt_fill

    for sheet in (summary_sheet, details_sheet):
        for column_index in range(1, sheet.max_column + 1):
            max_length = 0
            column_letter = get_column_letter(column_index)
            for row_index in range(1, sheet.max_row + 1):
                cell = sheet.cell(row=row_index, column=column_index)
                try:
                    value_length = len(str(cell.value)) if cell.value is not None else 0
                    max_length = max(max_length, value_length)
                except Exception:
                    pass
            sheet.column_dimensions[column_letter].width = min(max(max_length + 2, 12), 28)

    buffer = io.BytesIO()
    workbook.save(buffer)
    buffer.seek(0)
    return buffer


def render_attendance_report_docx(report_data):
    document = Document()

    title = document.add_heading("Attendance & Salary Report", level=0)
    title.alignment = 1

    meta = document.add_paragraph()
    meta.add_run(f"Scope: {report_data['scope_label']}\n").bold = True
    meta.add_run(f"Period: {report_data['start_date'].isoformat()} to {report_data['end_date'].isoformat()}\n")
    meta.add_run(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    if report_data["summary_rows"]:
        document.add_heading("Salary Summary", level=1)
        summary_table = document.add_table(rows=1, cols=6)
        summary_table.style = "Table Grid"
        headers = ["User ID", "Name", "Role", "Present Days", "Hours", "Net Salary"]
        for index, header in enumerate(headers):
            summary_table.rows[0].cells[index].text = header

        for row in report_data["summary_rows"]:
            cells = summary_table.add_row().cells
            cells[0].text = str(row["user_id"])
            cells[1].text = row["name"]
            cells[2].text = row["role"]
            cells[3].text = str(row["present_days"])
            cells[4].text = f"{row['total_hours']:.2f}"
            cells[5].text = f"{row['net_salary']:.2f}"

    document.add_heading("Attendance Sessions", level=1)
    if report_data["detail_rows"]:
        table = document.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        headers = ["User ID", "Name", "Date", "Day", "Check-In", "Check-Out", "Hours", "Status"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header

        for row in report_data["detail_rows"]:
            cells = table.add_row().cells
            cells[0].text = str(row["user_id"])
            cells[1].text = row["name"]
            cells[2].text = row["date"]
            cells[3].text = row["day"] or "-"
            cells[4].text = row["checkin_time"]
            cells[5].text = row["checkout_time"]
            cells[6].text = f"{row['worked_hours']:.2f}"
            cells[7].text = row["status"]
    else:
        document.add_paragraph("No attendance sessions found for the selected employee and date range.")

    if report_data["summary_rows"]:
        document.add_paragraph("")
        document.add_paragraph(
            "Estimated net salary: " + ", ".join(
                f"{row['name']} = {row['net_salary']:.2f}" for row in report_data["summary_rows"]
            )
        )

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def render_attendance_report_png(report_data):
    width = 1500
    summary_rows = report_data["summary_rows"]
    detail_rows = report_data["detail_rows"]
    row_height = 42
    header_height = 280
    table_height = 90 + max(len(detail_rows), 1) * row_height
    summary_height = 160 if summary_rows else 110
    height = header_height + summary_height + table_height + 60

    image = Image.new("RGB", (width, height), "#f4f7fb")
    draw = ImageDraw.Draw(image)
    title_font = ImageFont.load_default()
    body_font = ImageFont.load_default()

    def draw_text(x, y, text, fill="#0f1720"):
        draw.text((x, y), text, font=body_font, fill=fill)

    draw.rounded_rectangle((30, 24, width - 30, 140), radius=26, fill="#0d6efd")
    draw.text((60, 52), "Attendance & Salary Report", font=title_font, fill="white")
    draw.text((60, 82), f"Scope: {report_data['scope_label']}", font=body_font, fill="#dbeafe")
    draw.text((60, 104), f"Period: {report_data['start_date'].isoformat()} to {report_data['end_date'].isoformat()}", font=body_font, fill="#dbeafe")

    metrics = [
        ("Employees", str(len(summary_rows))),
        ("Sessions", str(len(detail_rows))),
        ("Total Hours", f"{sum(row['total_hours'] for row in summary_rows):.2f}" if summary_rows else "0.00"),
        ("Net Salary", f"{sum(row['net_salary'] for row in summary_rows):.2f}" if summary_rows else "0.00"),
    ]

    card_top = 170
    card_width = (width - 120) // 4
    for index, (label, value) in enumerate(metrics):
        x1 = 30 + index * card_width
        x2 = x1 + card_width - 12
        draw.rounded_rectangle((x1, card_top, x2, card_top + 72), radius=18, fill="#ffffff", outline="#d9e5f2", width=2)
        draw.text((x1 + 18, card_top + 14), label, font=body_font, fill="#5b6b79")
        draw.text((x1 + 18, card_top + 36), value, font=title_font, fill="#111827")

    summary_top = 270
    if summary_rows:
        draw.text((30, summary_top), "Salary Summary", font=title_font, fill="#0f1720")
        summary_top += 28
        header_y = summary_top
        summary_headers = [(30, "Employee"), (390, "Role"), (560, "Present"), (730, "Hours"), (880, "Net Salary")]
        draw.rounded_rectangle((30, header_y, width - 30, header_y + 38), radius=12, fill="#198754")
        for x, label in summary_headers:
            draw.text((x, header_y + 10), label, font=body_font, fill="white")
        summary_y = header_y + 44
        for index, row in enumerate(summary_rows[:10]):
            fill = "#ffffff" if index % 2 == 0 else "#eef5ff"
            draw.rounded_rectangle((30, summary_y, width - 30, summary_y + 36), radius=10, fill=fill, outline="#d9e5f2")
            draw.text((30, summary_y + 9), f"{row['name']} ({row['user_id']})", font=body_font, fill="#111827")
            draw.text((390, summary_y + 9), row["role"], font=body_font, fill="#111827")
            draw.text((560, summary_y + 9), str(row["present_days"]), font=body_font, fill="#111827")
            draw.text((730, summary_y + 9), f"{row['total_hours']:.2f}", font=body_font, fill="#111827")
            draw.text((880, summary_y + 9), f"{row['net_salary']:.2f}", font=body_font, fill="#111827")
            summary_y += 42
        if len(summary_rows) > 10:
            draw_text(30, summary_y + 2, "More employees available in PDF/Excel/Word exports.", fill="#5b6b79")
            summary_y += 24
    else:
        draw_text(30, summary_top, "No salary data found for the selected range.", fill="#5b6b79")
        summary_y = summary_top + 24

    table_top = summary_y + 34
    draw.text((30, table_top), "Attendance Sessions", font=title_font, fill="#0f1720")
    table_top += 28
    headers = [(30, "User"), (210, "Name"), (430, "Date"), (570, "In"), (710, "Out"), (850, "Hours"), (960, "Status")]
    draw.rounded_rectangle((30, table_top, width - 30, table_top + 38), radius=12, fill="#0d6efd")
    for x, label in headers:
        draw.text((x, table_top + 10), label, font=body_font, fill="white")
    table_top += 46

    max_rows = min(len(detail_rows), 25)
    for index, row in enumerate(detail_rows[:max_rows]):
        fill = "#ffffff" if index % 2 == 0 else "#f5fbf7"
        draw.rounded_rectangle((30, table_top, width - 30, table_top + 38), radius=10, fill=fill, outline="#d4e7dc")
        draw.text((30, table_top + 9), str(row["user_id"]), font=body_font, fill="#111827")
        draw.text((210, table_top + 9), row["name"], font=body_font, fill="#111827")
        draw.text((430, table_top + 9), row["date"], font=body_font, fill="#111827")
        draw.text((570, table_top + 9), row["checkin_time"], font=body_font, fill="#111827")
        draw.text((710, table_top + 9), row["checkout_time"], font=body_font, fill="#111827")
        draw.text((850, table_top + 9), f"{row['worked_hours']:.2f}", font=body_font, fill="#111827")
        draw.text((960, table_top + 9), row["status"], font=body_font, fill="#111827")
        table_top += 42

    if len(detail_rows) > max_rows:
        draw_text(30, table_top + 4, f"Showing the first {max_rows} sessions. Use PDF/Excel/Word for the complete export.", fill="#5b6b79")

    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def build_default_payroll_settings(db, user_id, month_key=None):
    if month_key is None:
        month_key = get_month_key()

    year, month, _, _ = get_month_bounds(month_key)
    expected_workdays = count_workdays_in_month(year, month)

    user_row = db.execute("""
        SELECT role
        FROM users
        WHERE id = ?
    """, (user_id,)).fetchone()
    role = ((user_row[0] if user_row else "employee") or "employee").strip().lower()

    standard_hours = 8.0
    # Default base pay policy
    monthly_salary = get_default_base_salary_for_role(role)
    hourly_rate = round(monthly_salary / max(expected_workdays * standard_hours, 1), 2)
    deduction_per_absent_day = round(monthly_salary / max(expected_workdays, 1), 2)

    return {
        "user_id": user_id,
        "monthly_salary": monthly_salary,
        "hourly_rate": hourly_rate,
        "workdays_per_month": expected_workdays,
        "standard_hours": standard_hours,
        "overtime_multiplier": 1.5,
        "deduction_per_absent_day": deduction_per_absent_day,
        "bonus_amount": 0.0,
        "updated_at": None,
        "is_default": True,
    }


def get_payroll_settings(db, user_id, month_key=None):
    row = db.execute("""
        SELECT user_id,
               monthly_salary,
               hourly_rate,
               workdays_per_month,
               standard_hours,
               overtime_multiplier,
               deduction_per_absent_day,
               bonus_amount,
               updated_at
        FROM payroll_settings
        WHERE user_id = ?
    """, (user_id,)).fetchone()

    if row:
        return {
            "user_id": row[0],
            "monthly_salary": row[1] or 0,
            "hourly_rate": row[2] or 0,
            "workdays_per_month": row[3] or 0,
            "standard_hours": row[4] or 8,
            "overtime_multiplier": row[5] or 1.5,
            "deduction_per_absent_day": row[6] or 0,
            "bonus_amount": row[7] or 0,
            "updated_at": row[8],
            "is_default": False,
        }

    return build_default_payroll_settings(db, user_id, month_key)


def build_payroll_settings_rows(db, month_key=None):
    employees = db.execute("""
        SELECT id, name, role
        FROM users
        ORDER BY role, name
    """).fetchall()

    settings_rows = []
    for user_id, name, role in employees:
        settings = get_payroll_settings(db, user_id, month_key)

        attendance_summary = db.execute("""
            SELECT COUNT(DISTINCT date),
                   ROUND(COALESCE(SUM(worked_hours), 0), 2)
            FROM attendance
            WHERE CAST(emp_id AS INTEGER) = ?
              AND checkout_time IS NOT NULL
              AND (? IS NULL OR substr(date, 1, 7) = ?)
        """, (user_id, month_key, month_key)).fetchone()

        settings_rows.append({
            "user_id": user_id,
            "name": name,
            "role": role,
            "settings": settings,
            "present_days": attendance_summary[0] or 0,
            "total_hours": attendance_summary[1] or 0,
        })

    return settings_rows


def upsert_payroll_settings(db, user_id, monthly_salary, hourly_rate, workdays_per_month,
                            standard_hours, overtime_multiplier, deduction_per_absent_day,
                            bonus_amount):
    db.execute("""
        INSERT INTO payroll_settings (
            user_id,
            monthly_salary,
            hourly_rate,
            workdays_per_month,
            standard_hours,
            overtime_multiplier,
            deduction_per_absent_day,
            bonus_amount,
            updated_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(user_id) DO UPDATE SET
            monthly_salary=excluded.monthly_salary,
            hourly_rate=excluded.hourly_rate,
            workdays_per_month=excluded.workdays_per_month,
            standard_hours=excluded.standard_hours,
            overtime_multiplier=excluded.overtime_multiplier,
            deduction_per_absent_day=excluded.deduction_per_absent_day,
            bonus_amount=excluded.bonus_amount,
            updated_at=excluded.updated_at
    """, (
        user_id,
        monthly_salary,
        hourly_rate,
        workdays_per_month,
        standard_hours,
        overtime_multiplier,
        deduction_per_absent_day,
        bonus_amount,
        datetime.datetime.now().isoformat(timespec="seconds"),
    ))


def build_payroll_snapshot(db, month_key):
    year, month, start_date, end_date = get_month_bounds(month_key)
    month_label = start_date.strftime("%B %Y")

    employees = db.execute("""
        SELECT id, name, role
        FROM users
        ORDER BY role, name
    """).fetchall()

    payroll_rows = []
    totals = {
        "expected_workdays": 0,
        "paid_days": 0,
        "present_days": 0,
        "absent_days": 0,
        "total_hours": 0.0,
        "overtime_hours": 0.0,
        "base_salary": 0.0,
        "overtime_pay": 0.0,
        "deductions": 0.0,
        "bonuses": 0.0,
        "gross_salary": 0.0,
        "net_salary": 0.0,
    }

    for employee in employees:
        user_id, name, role = employee
        settings = get_payroll_settings(db, user_id, month_key)
        expected_workdays = settings["workdays_per_month"] or count_workdays_in_month(year, month)
        completed_workdays = count_completed_workdays_in_month(year, month)

        attendance_rows = db.execute("""
            SELECT date,
                   ROUND(COALESCE(worked_hours, 0), 2)
            FROM attendance
            WHERE CAST(emp_id AS INTEGER) = ?
              AND substr(date, 1, 7) = ?
              AND checkout_time IS NOT NULL
        """, (user_id, month_key)).fetchall()

        worked_hours_by_day = {}
        for date_value, worked in attendance_rows:
            if not date_value:
                continue
            worked_hours_by_day[date_value] = worked_hours_by_day.get(date_value, 0.0) + float(worked or 0)

        total_hours = round(sum(worked_hours_by_day.values()), 2)
        present_days = len([d for d, hrs in worked_hours_by_day.items() if hrs > 0])

        standard_hours = settings["standard_hours"] or 8
        default_base_salary = get_default_base_salary_for_role(role)
        monthly_salary = max(float(settings["monthly_salary"] or 0), default_base_salary)
        hourly_rate = float(settings["hourly_rate"] or 0)
        overtime_multiplier = float(settings["overtime_multiplier"] or 1.5)
        bonus_amount = float(settings["bonus_amount"] or 0)
        _deduction_per_absent_day = float(settings["deduction_per_absent_day"] or 0)

        if monthly_salary > 0 and hourly_rate <= 0:
            effective_hourly_rate = monthly_salary / max(expected_workdays * standard_hours, 1)
        else:
            effective_hourly_rate = max(hourly_rate, 0)

        # Business rule: A day counts as a full paid day only when daily work >= 7 hours.
        paid_regular_hours = 0.0
        overtime_hours = 0.0
        full_day_credit = 0
        for day_hours in worked_hours_by_day.values():
            if day_hours >= MIN_FULL_DAY_HOURS:
                paid_regular_hours += standard_hours
                full_day_credit += 1
            else:
                paid_regular_hours += max(day_hours, 0.0)

            overtime_hours += max(day_hours - standard_hours, 0.0)

        # Absent should be month-to-date (completed workdays), not full-month estimate.
        absent_days = max(completed_workdays - present_days, 0)

        base_salary = round(max(paid_regular_hours, 0.0) * effective_hourly_rate, 2)
        overtime_pay = round(max(overtime_hours, 0.0) * effective_hourly_rate * overtime_multiplier, 2)
        deductions = 0.0
        gross_salary = round(base_salary + overtime_pay + bonus_amount, 2)
        net_salary = gross_salary

        payroll_row = {
            "user_id": user_id,
            "name": name,
            "role": role,
            "month": month_key,
            "month_label": month_label,
            "expected_workdays": expected_workdays,
            "completed_workdays": completed_workdays,
            "configured_base_salary": round(monthly_salary, 2),
            "paid_days": full_day_credit,
            "present_days": present_days,
            "absent_days": absent_days,
            "total_hours": total_hours,
            "overtime_hours": round(overtime_hours, 2),
            "base_salary": round(base_salary, 2),
            "overtime_pay": round(overtime_pay, 2),
            "deductions": round(deductions, 2),
            "bonuses": round(bonus_amount, 2),
            "gross_salary": round(gross_salary, 2),
            "net_salary": round(net_salary, 2),
            "settings": settings,
        }
        payroll_rows.append(payroll_row)

        totals["expected_workdays"] += expected_workdays
        totals["paid_days"] += full_day_credit
        totals["present_days"] += present_days
        totals["absent_days"] += absent_days
        totals["total_hours"] += total_hours
        totals["overtime_hours"] += payroll_row["overtime_hours"]
        totals["base_salary"] += payroll_row["base_salary"]
        totals["overtime_pay"] += payroll_row["overtime_pay"]
        totals["deductions"] += payroll_row["deductions"]
        totals["bonuses"] += payroll_row["bonuses"]
        totals["gross_salary"] += payroll_row["gross_salary"]
        totals["net_salary"] += payroll_row["net_salary"]

        db.execute("""
            INSERT INTO payroll_records (
                user_id,
                month,
                year,
                present_days,
                absent_days,
                expected_workdays,
                total_hours,
                overtime_hours,
                base_salary,
                overtime_pay,
                deductions,
                bonuses,
                gross_salary,
                net_salary,
                generated_at,
                status
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(user_id, month) DO UPDATE SET
                year=excluded.year,
                present_days=excluded.present_days,
                absent_days=excluded.absent_days,
                expected_workdays=excluded.expected_workdays,
                total_hours=excluded.total_hours,
                overtime_hours=excluded.overtime_hours,
                base_salary=excluded.base_salary,
                overtime_pay=excluded.overtime_pay,
                deductions=excluded.deductions,
                bonuses=excluded.bonuses,
                gross_salary=excluded.gross_salary,
                net_salary=excluded.net_salary,
                generated_at=excluded.generated_at,
                status=excluded.status
        """, (
            user_id,
            month_key,
            year,
            present_days,
            absent_days,
            expected_workdays,
            round(total_hours, 2),
            round(overtime_hours, 2),
            round(payroll_row["base_salary"], 2),
            round(payroll_row["overtime_pay"], 2),
            round(payroll_row["deductions"], 2),
            round(payroll_row["bonuses"], 2),
            round(payroll_row["gross_salary"], 2),
            round(payroll_row["net_salary"], 2),
            datetime.datetime.now().isoformat(timespec="seconds"),
            "generated",
        ))

    return payroll_rows, totals, month_label


ensure_payroll_schema()
ensure_leave_schema()


def parse_iso_date(date_str):
    return datetime.datetime.strptime(date_str, "%Y-%m-%d").date()


def apply_approved_leave_to_attendance(db, emp_id, start_date, end_date):
    """Replicate approved leave into attendance table as Leave Approved rows."""
    applied = 0
    skipped = 0
    current = start_date

    while current <= end_date:
        day_name = current.strftime("%A")
        date_str = current.isoformat()

        existing = db.execute("""
            SELECT id, checkin_time, checkout_time
            FROM attendance
            WHERE emp_id = ? AND date = ?
            ORDER BY id DESC
            LIMIT 1
        """, (str(emp_id), date_str)).fetchone()

        if existing:
            row_id, checkin_time, checkout_time = existing
            if checkin_time or checkout_time:
                skipped += 1
            else:
                db.execute("""
                    UPDATE attendance
                    SET day = ?, status = 'Leave Approved', worked_hours = 0
                    WHERE id = ?
                """, (day_name, row_id))
                applied += 1
        else:
            db.execute("""
                INSERT INTO attendance (emp_id, date, day, checkin_time, checkout_time, worked_hours, status)
                VALUES (?, ?, ?, NULL, NULL, 0, 'Leave Approved')
            """, (str(emp_id), date_str, day_name))
            applied += 1

        current += datetime.timedelta(days=1)

    return applied, skipped


def submit_leave_request(db, emp_id, subject, leave_type, start_date_str, end_date_str, reason):
    emp_id = (emp_id or "").strip()
    subject = (subject or "").strip()
    leave_type = (leave_type or "General").strip() or "General"
    start_date_str = (start_date_str or "").strip()
    end_date_str = (end_date_str or "").strip()
    reason = (reason or "").strip()

    if not emp_id or not subject or not start_date_str or not end_date_str or not reason:
        raise ValueError("All fields are required.")

    user_row = db.execute("SELECT id, name FROM users WHERE id = ?", (emp_id,)).fetchone()
    if not user_row:
        raise ValueError("Invalid employee ID. Please choose a valid user.")

    start_date = parse_iso_date(start_date_str)
    end_date = parse_iso_date(end_date_str)
    if end_date < start_date:
        raise ValueError("End date cannot be before start date.")

    total_days = (end_date - start_date).days + 1
    if total_days > 30:
        raise ValueError("Leave request cannot exceed 30 days in one request.")

    overlap = db.execute("""
        SELECT COUNT(*)
        FROM leave_requests
        WHERE emp_id = ?
          AND request_status IN ('Pending', 'Approved')
          AND NOT (end_date < ? OR start_date > ?)
    """, (emp_id, start_date_str, end_date_str)).fetchone()[0]

    if overlap:
        raise ValueError("An overlapping leave request already exists (pending/approved).")

    db.execute("""
        INSERT INTO leave_requests (
            emp_id, subject, leave_type, start_date, end_date, total_days, reason,
            request_status, applied_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, 'Pending', ?)
    """, (
        str(emp_id),
        subject,
        leave_type,
        start_date_str,
        end_date_str,
        total_days,
        reason,
        datetime.datetime.now().isoformat(timespec="seconds"),
    ))

    return total_days


def get_user_cache(force_refresh=False):
    global user_cache
    global user_cache_loaded_at

    now = time.time()
    if not force_refresh and user_cache and (now - user_cache_loaded_at) < 10.0:
        return user_cache

    db = get_db()
    rows = db.execute("SELECT id, name FROM users").fetchall()
    db.close()

    user_cache = {row[0]: row[1] for row in rows}
    user_cache_loaded_at = now
    return user_cache


def get_training_image_counts(force_refresh=False):
    global training_image_count_cache
    global training_image_count_loaded_at

    now = time.time()
    if not force_refresh and training_image_count_cache and (now - training_image_count_loaded_at) < 10.0:
        return training_image_count_cache

    counts = {}
    if os.path.exists("TrainingImage"):
        for filename in os.listdir("TrainingImage"):
            if not (filename.startswith("User.") and filename.endswith(".jpg")):
                continue

            parts = filename.split(".")
            if len(parts) < 4:
                continue

            try:
                uid = int(parts[1])
            except ValueError:
                continue

            counts[uid] = counts.get(uid, 0) + 1

    training_image_count_cache = counts
    training_image_count_loaded_at = now
    return training_image_count_cache


def clear_training_images_for_user(user_id):
    if not os.path.exists("TrainingImage"):
        return

    prefix = f"User.{user_id}."
    for filename in os.listdir("TrainingImage"):
        if filename.startswith(prefix):
            try:
                os.remove(os.path.join("TrainingImage", filename))
            except Exception:
                pass


def ensure_recognizer_is_current():
    """Reload recognizer if Trainner.yml changed on disk."""
    global recognizer
    global recognizer_model_mtime

    if not os.path.exists(MODEL_PATH):
        return False

    try:
        current_mtime = os.path.getmtime(MODEL_PATH)
    except Exception:
        return False

    if recognizer_model_mtime is not None and current_mtime == recognizer_model_mtime:
        return True

    try:
        fresh_recognizer = cv2.face.LBPHFaceRecognizer_create()
        fresh_recognizer.read(MODEL_PATH)
        recognizer = fresh_recognizer
        recognizer_model_mtime = current_mtime
        print("[MODEL] Reloaded recognizer from updated Trainner.yml")
        return True
    except Exception as e:
        print(f"[MODEL] Failed to reload recognizer: {e}")
        return False


def prepare_face_for_model(face_roi):
    if face_roi is None or face_roi.size == 0:
        return None

    try:
        processed = preprocess_face(face_roi)
        if processed is None or processed.size == 0:
            return None

        return resize_face_with_padding(
            processed,
            target_width=FACE_WIDTH,
            target_height=FACE_HEIGHT,
        )
    except Exception:
        return None


def build_temp_duplicate_recognizer(valid_user_ids):
    """Build a temporary recognizer from existing training images for duplicate checks."""
    if not os.path.exists("TrainingImage"):
        return None

    temp_faces = []
    temp_ids = []

    for filename in os.listdir("TrainingImage"):
        if not (filename.startswith("User.") and filename.endswith(".jpg")):
            continue

        parts = filename.split(".")
        if len(parts) < 4:
            continue

        try:
            uid = int(parts[1])
        except ValueError:
            continue

        if uid not in valid_user_ids:
            continue

        image_path = os.path.join("TrainingImage", filename)
        try:
            img = Image.open(image_path).convert("L")
            img_np = np.array(img, dtype=np.uint8)
            if img_np.size < 100:
                continue

            if img_np.shape == (FACE_HEIGHT, FACE_WIDTH):
                prepared = img_np
            else:
                prepared = prepare_face_for_model(img_np)
            if prepared is None:
                continue

            temp_faces.append(prepared)
            temp_ids.append(uid)
        except Exception:
            continue

    if len(temp_faces) < MIN_IMAGES_PER_USER:
        return None

    if len(set(temp_ids)) < 1:
        return None

    try:
        temp_recognizer = cv2.face.LBPHFaceRecognizer_create()
        temp_recognizer.train(temp_faces, np.array(temp_ids, dtype=np.int32))
        return temp_recognizer
    except Exception:
        return None


def detect_duplicate_enrollment_face(captured_faces, enrolling_user_id):
    """
    Detect duplicate enrollment by requiring stable multi-frame consensus.
    Only considers DB-backed users that have enough training images.
    """
    if not captured_faces:
        return None, None, None

    users = get_user_cache(force_refresh=True)
    sample_counts = get_training_image_counts(force_refresh=True)
    valid_user_ids = {
        uid for uid in users.keys()
        if uid != enrolling_user_id and sample_counts.get(uid, 0) >= DUPLICATE_MIN_IMAGES_PER_USER
    }

    if not valid_user_ids:
        return None, None, {
            "duplicate_conf_threshold": DUPLICATE_CONF_THRESHOLD,
            "strong_duplicate_conf": STRONG_DUPLICATE_CONF,
            "required_hits": DUPLICATE_MIN_HITS,
            "confidence_margin": 0,
            "top_candidates": []
        }

    votes = Counter()
    conf_by_user = {}

    predictor = None
    if os.path.exists("TrainingImageLabel/Trainner.yml"):
        try:
            recognizer.read("TrainingImageLabel/Trainner.yml")
            predictor = recognizer
        except Exception:
            predictor = None

    # Fallback when model is missing or unreadable: train a temporary recognizer from existing images.
    if predictor is None:
        predictor = build_temp_duplicate_recognizer(valid_user_ids)

    if predictor is None:
        return None, None, {
            "duplicate_conf_threshold": DUPLICATE_CONF_THRESHOLD,
            "strong_duplicate_conf": STRONG_DUPLICATE_CONF,
            "required_hits": DUPLICATE_MIN_HITS,
            "confidence_margin": 0,
            "top_candidates": [],
            "note": "duplicate_predictor_unavailable"
        }

    for face_roi in captured_faces:
        # captured_faces already stores normalized 200x200 grayscale frames from enrollment.
        # Avoid double preprocessing here because it can collapse identity-specific details.
        if (
            isinstance(face_roi, np.ndarray)
            and len(face_roi.shape) == 2
            and face_roi.shape == (FACE_HEIGHT, FACE_WIDTH)
        ):
            prepared = face_roi
        else:
            prepared = prepare_face_for_model(face_roi)
            if prepared is None:
                continue

        try:
            pred_user_id, conf = predictor.predict(prepared)
        except Exception:
            continue

        if pred_user_id in valid_user_ids and conf <= DUPLICATE_CONF_THRESHOLD:
            votes[pred_user_id] += 1
            conf_by_user.setdefault(pred_user_id, []).append(float(conf))

    required_hits = max(DUPLICATE_MIN_HITS, int(math.ceil(len(captured_faces) * DUPLICATE_REQUIRED_RATIO)))
    top_candidates = []

    for uid, hit_count in votes.items():
        conf_list = conf_by_user.get(uid, [])
        if not conf_list:
            continue

        top_candidates.append({
            "user_id": uid,
            "name": users.get(uid, f"User {uid}"),
            "hits": hit_count,
            "best_conf": round(min(conf_list), 2),
            "avg_conf": round(sum(conf_list) / len(conf_list), 2),
        })

    top_candidates.sort(key=lambda row: (-row["hits"], row["avg_conf"]))

    debug_payload = {
        "duplicate_conf_threshold": DUPLICATE_CONF_THRESHOLD,
        "strong_duplicate_conf": STRONG_DUPLICATE_CONF,
        "required_hits": required_hits,
        "required_conf_margin": DUPLICATE_MIN_CONF_MARGIN,
        "confidence_margin": 0,
        "top_candidates": top_candidates[:3],
    }

    if len(top_candidates) >= 2:
        best_avg = top_candidates[0]["avg_conf"]
        second_avg = top_candidates[1]["avg_conf"]
        debug_payload["best_avg_conf"] = round(best_avg, 2)
        debug_payload["second_best_avg_conf"] = round(second_avg, 2)
        debug_payload["confidence_margin"] = round(max(0.0, second_avg - best_avg), 2)

    if not top_candidates:
        return None, None, debug_payload

    winner = top_candidates[0]
    winner_hits = winner["hits"]
    winner_best_conf = winner["best_conf"]
    winner_avg_conf = winner["avg_conf"]

    # Require either strong confidence with enough hits, or clear consensus by hit count.
    is_strong_duplicate = winner_best_conf <= STRONG_DUPLICATE_CONF and winner_hits >= required_hits
    is_consensus_duplicate = winner_hits >= (required_hits + 1) and winner_avg_conf <= DUPLICATE_CONF_THRESHOLD

    # If multiple candidates are close, treat as ambiguous and allow enrollment.
    if len(top_candidates) >= 2:
        confidence_margin = debug_payload.get("confidence_margin", 0)
        if confidence_margin < DUPLICATE_MIN_CONF_MARGIN:
            return None, None, debug_payload

    # When only one candidate appears, require even stronger evidence to avoid false duplicate blocks.
    if len(top_candidates) == 1:
        is_consensus_duplicate = winner_hits >= (required_hits + 2) and winner_avg_conf <= STRONG_DUPLICATE_CONF

    if is_strong_duplicate or is_consensus_duplicate:
        return winner["user_id"], winner["name"], debug_payload

    return None, None, debug_payload


# ============================================================================
# COMPLEX VERIFICATION FUNCTIONS REMOVED
# ============================================================================
# The following verification functions (histogram correlation, LBP comparison,
# iris detection, mask detection) were removed in favor of a simplified,
# more reliable recognition pipeline that uses:
# 1. Direct LBPH confidence scoring
# 2. Stable frame voting across multiple frames
# 3. Clean, maintainable code
#
# Original functions removed:
#   - refresh_face_samples_cache()
#   - verify_prediction_with_samples()
#   - verify_prediction_with_iris()
#   - verify_prediction_with_upper_samples()
#   - detect_probable_mask()
#   - _extract_iris_signature()
#   - _extract_upper_face()
#   - _edge_density()
#   - _gray_hist()
#   - _lbp_hist()
#   - _orb_desc()
#   - _orb_mean_distance()
#   - infer_user_id_from_face_image()
#
# If you need to restore these functions, check git history or
# contact the development team.
# ============================================================================

# ------------------ Attendance Logic ------------------
def mark_attendance(user_id, face_image):
    global attendance_mode

    print(f"\n{'='*60}")
    print(f"[MARK_ATTENDANCE] Function called with:")
    print(f"  user_id parameter: {user_id} (type: {type(user_id)})")
    print(f"  attendance_mode: {attendance_mode}")
    print(f"{'='*60}")

    now = datetime.datetime.now()
    today = now.strftime("%Y-%m-%d")
    day = now.strftime("%A")
    time_now_file = now.strftime("%H-%M-%S")
    time_now_db = now.strftime("%H:%M:%S")

    db = get_db()
    cur = db.cursor()

    # CRITICAL FIX: VALIDATE USER EXISTS
    user_check = cur.execute("""
        SELECT id, name FROM users WHERE id = ?
    """, (user_id,)).fetchone()
    
    if not user_check:
        get_user_cache(force_refresh=True)
        db.close()
        print(f"❌ [ERROR] Recognized user_id {user_id} does NOT exist in database!")
        print(f"   Valid user IDs in DB:")
        db = get_db()
        valid_ids = db.execute("SELECT id, name FROM users").fetchall()
        for vid, vname in valid_ids:
            print(f"      - ID {vid}: {vname}")
        db.close()
        return f"invalid_user_{user_id}"
    
    user_name = user_check[1]
    print(f"✅ User validation successful: ID={user_id}, Name={user_name}")

    os.makedirs("static/attendance_images", exist_ok=True)

    # CHECK-IN
    if attendance_mode == "checkin":
        cur.execute("""
            SELECT id FROM attendance
            WHERE emp_id=? AND date=? AND checkout_time IS NULL
            ORDER BY id DESC LIMIT 1
        """, (user_id, today))

        open_session = cur.fetchone()

        if open_session:
            db.close()
            print(f"⚠️  User {user_id} already checked in today")
            return "already_checked_in"

        filename = f"checkin_{user_id}_{today}_{time_now_file}.jpg"
        full_path = os.path.join("static/attendance_images", filename)
        cv2.imwrite(full_path, face_image)
        db_path = f"attendance_images/{filename}"

        cur.execute("""
            INSERT INTO attendance 
            (emp_id, date, day, checkin_time, status, checkin_image)
            VALUES (?, ?, ?, ?, 'Present', ?)
        """, (user_id, today, day, time_now_db, db_path))

        db.commit()
        db.close()
        print(f"✅ Check-in recorded for User {user_id} ({user_name})")
        print(f"   Date: {today}, Time: {time_now_db}")
        return f"checkin_success_{user_name}"

    # CHECK-OUT
    elif attendance_mode == "checkout":
        cur.execute("""
            SELECT id, checkin_time, checkin_image
            FROM attendance
            WHERE emp_id=? AND date=? AND checkout_time IS NULL
            ORDER BY id DESC LIMIT 1
        """, (user_id, today))

        open_session = cur.fetchone()

        if not open_session:
            db.close()
            print(f"⚠️  No open session for User {user_id} today")
            return "already_checked_out"

        record_id, checkin_time, checkin_image_path = open_session

        # Strong guard: ensure checkout face matches this session's checkin face.
        try:
            current_prepared = prepare_face_for_model(face_image)
            checkin_prepared = None

            if checkin_image_path:
                checkin_full_path = os.path.join("static", checkin_image_path)
                if os.path.exists(checkin_full_path):
                    checkin_img = cv2.imread(checkin_full_path, cv2.IMREAD_GRAYSCALE)
                    if checkin_img is not None:
                        checkin_prepared = prepare_face_for_model(checkin_img)

            if current_prepared is not None and checkin_prepared is not None:
                identity_score = face_similarity(current_prepared, checkin_prepared)
                if identity_score is not None and identity_score < 0.52:
                    db.close()
                    print(
                        f"❌ [CHECKOUT BLOCKED] Face mismatch for user {user_id}. "
                        f"checkin-vs-checkout similarity={identity_score:.3f}"
                    )
                    return "checkout_face_mismatch"
        except Exception as e:
            print(f"⚠️ Checkout identity verification skipped due to error: {e}")

        login_time = datetime.datetime.strptime(checkin_time, "%H:%M:%S")
        logout_time = datetime.datetime.strptime(time_now_db, "%H:%M:%S")
        worked_hours = (logout_time - login_time).total_seconds() / 3600

        filename = f"checkout_{user_id}_{today}_{time_now_file}.jpg"
        full_path = os.path.join("static/attendance_images", filename)
        cv2.imwrite(full_path, face_image)
        db_path = f"attendance_images/{filename}"

        cur.execute("""
            UPDATE attendance
            SET checkout_time=?, worked_hours=?, checkout_image=?
            WHERE id=?
        """, (time_now_db, worked_hours, db_path, record_id))

        db.commit()
        db.close()
        print(f"✅ Check-out recorded for User {user_id} ({user_name})")
        print(f"   Date: {today}, Time: {time_now_db}, Worked: {worked_hours:.2f} hrs")
        return f"checkout_success_{user_name}"
    
    db.close()
    return "error_invalid_mode"


def get_attendance_stats():
    db = get_db()

    today_count = db.execute("""
        SELECT COUNT(DISTINCT user_id)
        FROM attendance
        WHERE date = date('now')
    """).fetchone()[0]

    total_count = db.execute("""
        SELECT COUNT(*)
        FROM attendance
    """).fetchone()[0]

    db.close()

    percentage = 0
    if TOTAL_EMPLOYEES > 0:
        percentage = round((today_count / TOTAL_EMPLOYEES) * 100, 2)

    return today_count, total_count, percentage

# ------------------ Video Streaming ------------------
def generate_frames():
    """
    Simplified face detection and recognition pipeline.
    - Detects faces in each frame
    - Recognizes using LBPH (Local Binary Patterns Histograms)
    - Requires stable matches across several frames before recording attendance
    - Clean, maintainable, and reliable
    """
    global attendance_status, success_recognition, failed_recognition
    global camera, last_attendance_time

    if face_cascade is None:
        print("❌ ERROR: Face cascade classifier not loaded!")
        return

    recognition_history = deque(maxlen=max(REQUIRED_STABLE_MATCHES + 4, 8))
    frames_without_detection = 0
    
    # Initialize camera
    if camera is None:
        try:
            if os.name == "nt":
                camera = cv2.VideoCapture(CAMERA_INDEX, cv2.CAP_DSHOW)
            else:
                camera = cv2.VideoCapture(CAMERA_INDEX)
            
            camera.set(cv2.CAP_PROP_BUFFERSIZE, 1)
            camera.set(cv2.CAP_PROP_FRAME_WIDTH, 640)
            camera.set(cv2.CAP_PROP_FRAME_HEIGHT, 480)
            print("✅ Camera initialized successfully")
        except Exception as e:
            print(f"❌ Camera initialization error: {e}")
            return

    frame_count = 0

    while camera_enabled:
        try:
            loop_start = time.time()
            
            ret, frame = camera.read()
            if not ret or frame is None:
                print("⚠️ Camera read failed")
                break
            
            frame_count += 1

            if frame_count % 15 == 1:
                ensure_recognizer_is_current()

            confidence = None
            status_text = "Detecting Face..."
            status_subtext = "Keep one face centered"
            status_color = (0, 200, 255)  # Blue
            box_color = (0, 100, 255)

            # Detect faces
            face_coords = detect_faces_optimized(frame, face_cascade)
            
            if len(face_coords) == 0:
                frames_without_detection += 1
                status_text = "No face detected"
                status_subtext = "Position your face in the frame"
                status_color = (200, 100, 100)  # Light blue
                recognition_history.clear()
            
            elif len(face_coords) > 1:
                # Multiple faces - ask user to show only one
                status_text = "Multiple Faces Detected"
                status_subtext = "Please show only one face"
                status_color = (0, 0, 255)  # Red
                box_color = (0, 0, 255)
                recognition_history.clear()
                
                for (x, y, w, h) in face_coords:
                    frame = draw_face_rectangle(frame, (x, y, w, h), "Multiple", color=box_color)
            
            else:
                # Single face detected
                (x, y, w, h) = face_coords[0]
                frames_without_detection = 0
                
                # Extract and prepare face for recognition
                face_prepared = extract_and_prepare_face(frame, (x, y, w, h), face_cascade)
                
                if face_prepared is None or face_prepared.size == 0:
                    status_text = "Face too small or unclear"
                    status_subtext = "Move closer to camera"
                    status_color = (0, 100, 200)
                    frame = draw_face_rectangle(frame, (x, y, w, h), "Unclear", color=(0, 100, 200))
                
                else:
                    # Check if trained model is available
                    if not os.path.exists("TrainingImageLabel/Trainner.yml"):
                        status_text = "Model Not Trained"
                        status_subtext = "Please enroll users and train first"
                        status_color = (0, 140, 255)
                        frame = draw_face_rectangle(frame, (x, y, w, h), "No Model", color=(0, 140, 255))
                    
                    else:
                        try:
                            # Recognize face using both original and flipped orientation.
                            user_id, confidence, prediction_debug = predict_face_with_flip(
                                recognizer,
                                face_prepared
                            )
                            if user_id is None or confidence is None:
                                prediction_error = prediction_debug.get("error", "unknown")
                                if prediction_error in {"ambiguous_flip_prediction", "low_confidence_flip_prediction"}:
                                    status_text = "Low Confidence Match"
                                    if prediction_error == "ambiguous_flip_prediction":
                                        status_subtext = "Face orientation is unstable - hold still and face forward"
                                    else:
                                        status_subtext = "Face not clear enough - move closer and look straight"
                                    status_color = (0, 165, 255)
                                    box_color = (0, 165, 255)
                                    recognition_history.clear()
                                    failed_recognition += 1
                                    label = status_text
                                    frame = draw_face_rectangle(
                                        frame, (x, y, w, h), label,
                                        confidence=None,
                                        color=box_color
                                    )
                                    continue
                                raise ValueError(f"Prediction failed: {prediction_error}")
                            
                            # Get user info
                            users = get_user_cache()
                            sample_counts = get_training_image_counts()
                            gallery_ok, gallery_debug = verify_prediction_with_gallery(
                                face_prepared,
                                user_id,
                            )

                            # Strict mode: never remap to another ID. Accept only the predicted ID.
                            if user_id in users:
                                user_name = users[user_id]
                                effective_threshold = min(THRESHOLD, 46)
                                enough_user_samples = sample_counts.get(user_id, 0) >= MIN_IMAGES_PER_USER
                                is_consistent_prediction = prediction_debug.get("agreement", False)
                                flip_resolved_by_confidence = bool(prediction_debug.get("flip_resolved_by_confidence", False))

                                gallery_top_user = None
                                gallery_margin = 0.0
                                gallery_best = 0.0
                                gallery_avg_top = 0.0
                                gallery_support = 0
                                if isinstance(gallery_debug, dict):
                                    gallery_top_user = gallery_debug.get("top_user_id")
                                    gallery_margin = float(gallery_debug.get("candidate_margin") or 0.0)
                                    gallery_best = float(gallery_debug.get("candidate_best_score") or 0.0)
                                    gallery_avg_top = float(gallery_debug.get("candidate_avg_top") or 0.0)
                                    gallery_support = int(gallery_debug.get("candidate_supporting_samples") or 0)

                                strong_gallery_match = (
                                    gallery_ok
                                    and gallery_top_user == user_id
                                    and gallery_best >= 0.50
                                    and gallery_avg_top >= 0.43
                                    and gallery_support >= 5
                                    and gallery_margin >= 0.08
                                )

                                borderline_confidence = confidence > effective_threshold and confidence <= 55
                                borderline_gallery_match = (
                                    gallery_ok
                                    and gallery_top_user == user_id
                                    and gallery_best >= 0.54
                                    and gallery_avg_top >= 0.45
                                    and gallery_support >= 5
                                    and gallery_margin >= 0.06
                                )

                                consistency_gate_ok = (
                                    is_consistent_prediction
                                    or (
                                        flip_resolved_by_confidence
                                        and strong_gallery_match
                                        and confidence <= effective_threshold
                                    )
                                )

                                frame_passes_identity = (
                                    confidence <= 55
                                    and enough_user_samples
                                    and consistency_gate_ok
                                    and (strong_gallery_match or borderline_gallery_match)
                                )
                                
                                if frame_passes_identity:
                                    # Good frame-level match. Show name only after stable lock.
                                    status_color = (0, 255, 0)  # Green
                                    box_color = (0, 255, 0)
                                    
                                    # Add to recognition history
                                    recognition_history.append({
                                        "user_id": user_id,
                                        "user_name": user_name,
                                        "confidence": confidence,
                                        "timestamp": time.time(),
                                        "face_image": frame[y:y+h, x:x+w].copy()
                                    })

                                    same_person_count = sum(
                                        1 for r in recognition_history
                                        if r["user_id"] == user_id
                                    )

                                    required_matches = REQUIRED_STABLE_MATCHES + 3 if borderline_confidence else REQUIRED_STABLE_MATCHES

                                    if same_person_count >= required_matches:
                                        status_text = f"Verified: {user_name}"
                                        status_subtext = "Recording attendance..."
                                    else:
                                        status_text = "Verifying Identity"
                                        status_subtext = f"Hold still ({same_person_count}/{required_matches})"
                                    
                                    # Check if we have enough stable matches
                                    if len(recognition_history) >= required_matches:
                                        # Require all recent frames to agree on the same identity.
                                        matching_confidences = [
                                            r["confidence"]
                                            for r in recognition_history
                                            if r["user_id"] == user_id
                                        ]
                                        avg_confidence = (
                                            sum(matching_confidences) / len(matching_confidences)
                                            if matching_confidences else 999.0
                                        )

                                        stable_accept = (
                                            same_person_count >= required_matches
                                            and avg_confidence <= (effective_threshold - 2)
                                        )

                                        borderline_accept = (
                                            borderline_confidence
                                            and same_person_count >= required_matches
                                            and avg_confidence <= 53
                                            and gallery_best >= 0.56
                                            and gallery_avg_top >= 0.47
                                            and gallery_support >= 6
                                            and gallery_margin >= 0.08
                                        )

                                        if stable_accept or borderline_accept:
                                            # Record attendance
                                            current_time = time.time()
                                            if current_time - last_attendance_time > ATTENDANCE_COOLDOWN:
                                                best_frame_idx = min(
                                                    range(len(recognition_history)),
                                                    key=lambda i: recognition_history[i]["confidence"]
                                                )
                                                face_image = recognition_history[best_frame_idx]["face_image"]
                                                
                                                result = mark_attendance(user_id, face_image)
                                                attendance_status = result
                                                success_recognition += 1
                                                last_attendance_time = current_time
                                                recognition_history.clear()
                                                
                                                print(f"✅ [ATTENDANCE] {user_name} (ID: {user_id}) - {result}")
                                else:
                                    # Weak match
                                    status_color = (0, 165, 255)  # Orange
                                    box_color = (0, 165, 255)
                                    status_text = "Low Confidence Match"
                                    if confidence > 55:
                                        status_subtext = "Face not clear enough - move closer and look straight"
                                    elif not enough_user_samples:
                                        status_subtext = "Insufficient training samples for this user"
                                    elif not consistency_gate_ok:
                                        status_subtext = "Unstable face match - hold still and face forward"
                                    elif borderline_confidence and gallery_top_user == user_id:
                                        status_subtext = "Need a few stable frames to confirm identity"
                                    elif not strong_gallery_match:
                                        status_subtext = "Face-ID mismatch blocked for safety"
                                    else:
                                        status_subtext = "Look straight, better lighting"
                                    recognition_history.clear()
                                    failed_recognition += 1
                            
                            else:
                                # Unknown user ID
                                status_color = (255, 100, 100)  # Light blue
                                box_color = (255, 100, 100)
                                status_text = "Unknown Person"
                                status_subtext = f"User ID: {user_id} not found"
                                recognition_history.clear()
                        
                        except Exception as e:
                            print(f"⚠️ Recognition error: {e}")
                            status_text = "Recognition Error"
                            status_subtext = "Try again"
                            status_color = (0, 0, 255)
                            recognition_history.clear()
                
                # Draw face rectangle with label
                label = status_text.split(":")[1].strip() if ":" in status_text else status_text
                frame = draw_face_rectangle(
                    frame, (x, y, w, h), label,
                    confidence=confidence if 'confidence' in locals() else None,
                    color=box_color
                )

            # Draw status banner at top
            h_frame, w_frame = frame.shape[:2]
            
            # Semi-transparent background
            overlay = frame.copy()
            cv2.rectangle(overlay, (0, 0), (w_frame, 100), (20, 20, 30), -1)
            frame = cv2.addWeighted(overlay, 0.7, frame, 0.3, 0)
            
            # Status text
            cv2.putText(
                frame, status_text,
                (20, 40),
                cv2.FONT_HERSHEY_SIMPLEX,
                1.0, status_color, 2
            )
            
            cv2.putText(
                frame, status_subtext,
                (20, 70),
                cv2.FONT_HERSHEY_SIMPLEX,
                0.6, (200, 200, 200), 1
            )
            
            # Encode frame
            _, buffer = cv2.imencode(".jpg", frame)
            
            yield (b"--frame\r\n"
                   b"Content-Type: image/jpeg\r\n\r\n" +
                   buffer.tobytes() + b"\r\n")
            
            # Frame rate control
            loop_time = time.time() - loop_start
            sleep_time = (1.0 / TARGET_STREAM_FPS) - loop_time
            if sleep_time > 0:
                time.sleep(sleep_time)
        
        except Exception as e:
            print(f"⚠️ Frame processing error: {e}")
            continue

    # Cleanup
    if camera is not None:
        camera.release()
        camera = None
    
    print("✓ Video stream ended")



        
def train_model():
    """Train face recognition model - uses pre-detected face ROIs with consistent sizing"""
    recognizer = cv2.face.LBPHFaceRecognizer_create()
    
    # CRITICAL: All images must be resized to same dimensions for LBPH training
    FACE_WIDTH = 200
    FACE_HEIGHT = 200

    faces = []
    ids = []
    min_required_samples = MIN_IMAGES_PER_USER

    if not os.path.exists("TrainingImage"):
        os.makedirs("TrainingImage")
        print("ERROR: TrainingImage folder was empty. Please enroll users first.")
        return False

    image_files = [f for f in os.listdir("TrainingImage") if f.endswith(".jpg")]
    if len(image_files) == 0:
        print("ERROR: No training images found. Please enroll users first.")
        return False

    print(f"[TRAINING] Processing {len(image_files)} images...")

    db = get_db()
    valid_user_ids = set(row[0] for row in db.execute("SELECT id FROM users").fetchall())
    db.close()

    image_paths = [os.path.join("TrainingImage", f) for f in image_files]
    user_face_count = {}

    for image_path in image_paths:
        try:
            # Load image as grayscale
            gray_img = Image.open(image_path).convert("L")
            img_np = np.array(gray_img, "uint8")

            # Verify minimal size
            if img_np.size < 100:
                continue

            filename = os.path.split(image_path)[-1]
            parts = filename.split(".")
            if len(parts) < 4 or parts[0] != "User":
                continue

            try:
                user_id = int(parts[1])
            except ValueError:
                continue

            if user_id not in valid_user_ids:
                continue

            # Enrollment already stores normalized 200x200 grayscale faces.
            if img_np.shape == (FACE_HEIGHT, FACE_WIDTH):
                resized_img = img_np
            else:
                resized_img = prepare_face_for_model(img_np)
            if resized_img is None:
                continue
            
            faces.append(resized_img)
            ids.append(user_id)
            # Add a mirrored copy so prediction remains stable even if camera feed is flipped.
            faces.append(cv2.flip(resized_img, 1))
            ids.append(user_id)
            
            user_face_count[user_id] = user_face_count.get(user_id, 0) + 1
                
        except Exception:
            continue

    if len(faces) == 0:
        print(f"❌ ERROR: No valid faces found for training!")
        return False

    # Avoid training identities that have too few samples (high mismatch risk).
    eligible_user_ids = {uid for uid, count in user_face_count.items() if count >= min_required_samples}
    if not eligible_user_ids:
        print(f"❌ ERROR: No users have enough samples. Need at least {min_required_samples} images per user.")
        return False

    filtered_faces = []
    filtered_ids = []
    for face_img, uid in zip(faces, ids):
        if uid in eligible_user_ids:
            filtered_faces.append(face_img)
            filtered_ids.append(uid)

    skipped_users = {uid: cnt for uid, cnt in user_face_count.items() if uid not in eligible_user_ids}
    if skipped_users:
        print(f"⚠️ Skipping users with insufficient samples (<{min_required_samples}): {skipped_users}")

    faces = filtered_faces
    ids = filtered_ids

    if len(faces) == 0:
        print("❌ ERROR: No eligible training samples after filtering.")
        return False

    print(f"[TRAINING] Training model with {len(faces)} samples from {len(eligible_user_ids)} users...")

    try:
        recognizer.train(faces, np.array(ids, dtype=np.int32))
        
        os.makedirs("TrainingImageLabel", exist_ok=True)
        recognizer.save("TrainingImageLabel/Trainner.yml")
        
        print(f"✅ Training complete!")
        return True
    except Exception as e:
        print(f"❌ Training error: {str(e)}")
        return False


def reload_recognizer_from_disk():
    """Reload latest trained model into memory."""
    global recognizer
    global recognizer_model_mtime

    if not os.path.exists(MODEL_PATH):
        raise FileNotFoundError("Trained model not found on disk.")

    new_recognizer = cv2.face.LBPHFaceRecognizer_create()
    new_recognizer.read(MODEL_PATH)
    recognizer = new_recognizer
    recognizer_model_mtime = os.path.getmtime(MODEL_PATH)
    # Note: Complex face samples cache functionality removed for simplicity


class TrainingJobManager:
    """Simple in-process background training job manager."""

    def __init__(self):
        self._lock = threading.Lock()
        self._jobs = {}
        self._active_job_id = None

    def start_job(self, include_smoke_test=False):
        with self._lock:
            if self._active_job_id:
                active = self._jobs.get(self._active_job_id, {})
                if active.get("status") in ("queued", "running"):
                    return self._active_job_id

            job_id = str(uuid.uuid4())
            self._jobs[job_id] = {
                "id": job_id,
                "status": "queued",
                "progress": 0,
                "stage": "queued",
                "message": "Training job queued",
                "include_smoke_test": include_smoke_test,
                "created_at": time.time(),
                "started_at": None,
                "finished_at": None,
                "error": None,
            }
            self._active_job_id = job_id

        worker = threading.Thread(
            target=self._run_job,
            args=(job_id,),
            daemon=True
        )
        worker.start()
        return job_id

    def get_job(self, job_id):
        with self._lock:
            job = self._jobs.get(job_id)
            return dict(job) if job else None

    def get_active_job(self):
        with self._lock:
            if not self._active_job_id:
                return None
            active = self._jobs.get(self._active_job_id)
            return dict(active) if active else None

    def _update_job(self, job_id, **updates):
        with self._lock:
            if job_id in self._jobs:
                self._jobs[job_id].update(updates)

    def _run_job(self, job_id):
        try:
            self._update_job(
                job_id,
                status="running",
                stage="precheck",
                progress=10,
                message="Validating training dataset",
                started_at=time.time(),
            )

            success = train_model()
            if not success:
                raise RuntimeError("Training failed. Check image quality and sample count per user.")

            self._update_job(
                job_id,
                stage="model_reload",
                progress=80,
                message="Reloading recognition model",
            )
            reload_recognizer_from_disk()

            job = self.get_job(job_id) or {}
            if job.get("include_smoke_test"):
                self._update_job(
                    job_id,
                    stage="smoke_test",
                    progress=92,
                    message="Running model smoke test",
                )
                if not os.path.exists("TrainingImageLabel/Trainner.yml"):
                    raise RuntimeError("Model smoke test failed: model file missing.")

            self._update_job(
                job_id,
                status="success",
                stage="completed",
                progress=100,
                message="Training completed successfully",
                finished_at=time.time(),
            )
        except Exception as e:
            self._update_job(
                job_id,
                status="fail",
                stage="failed",
                progress=100,
                message="Training failed",
                error=str(e),
                finished_at=time.time(),
            )
        finally:
            with self._lock:
                if self._active_job_id == job_id:
                    self._active_job_id = None


training_jobs = TrainingJobManager()
app.register_blueprint(create_training_blueprint(training_jobs))

def eye_aspect_ratio(landmarks, left_eye, right_eye):
    def distance(p1, p2):
        return math.dist((p1.x, p1.y), (p2.x, p2.y))

    left_ear = (
        distance(landmarks[left_eye[1]], landmarks[left_eye[5]]) +
        distance(landmarks[left_eye[2]], landmarks[left_eye[4]])
    ) / (2.0 * distance(landmarks[left_eye[0]], landmarks[left_eye[3]]))

    right_ear = (
        distance(landmarks[right_eye[1]], landmarks[right_eye[5]]) +
        distance(landmarks[right_eye[2]], landmarks[right_eye[4]])
    ) / (2.0 * distance(landmarks[right_eye[0]], landmarks[right_eye[3]]))

    return (left_ear + right_ear) / 2.0

def admin_required():
    return session.get("admin", False)

def admin():
    if not admin_required():
        return redirect("/login")
    return render_template("admin.html")





# ------------------ Routes ------------------
@app.route("/")
def index():
    return render_template("index.html")

@app.route("/video")
def video():
    return Response(
        generate_frames(),
        mimetype="multipart/x-mixed-replace; boundary=frame"
    )

# ---------------- COMBINED ATTENDANCE ----------------
@app.route("/attendance")
def attendance():

    db = get_db()

    # Fetch all attendance records with user role
    rows_raw = db.execute("""
        SELECT attendance.id,
               users.id,
               users.name,
               attendance.date,
               attendance.day,
               attendance.checkin_time,
               attendance.checkout_time,
               attendance.worked_hours,
             attendance.status,
               attendance.checkin_image,
               attendance.checkout_image,
               users.role
        FROM attendance
        INNER JOIN users 
            ON CAST(attendance.emp_id AS INTEGER) = users.id
        ORDER BY attendance.date DESC, attendance.checkin_time DESC
    """).fetchall()

    users_map = {
        r[0]: {"name": r[1], "role": r[2]}
        for r in db.execute("SELECT id, name, role FROM users").fetchall()
    }

    corrected_rows = []
    for row in rows_raw:
        row_list = list(row)
        # Note: Image-based inference functionality removed for simplicity
        # The emp_id in attendance records is the source of truth
        corrected_rows.append(tuple(row_list))

    # Total users (all roles)
    users = db.execute("""
        SELECT id, name, role
        FROM users
        ORDER BY name ASC
    """).fetchall()

    total_users = db.execute("""
        SELECT COUNT(*) FROM users
    """).fetchone()[0]

    today = datetime.date.today().isoformat()

    # Today's attendance count (all roles)
    today_count = db.execute("""
        SELECT COUNT(DISTINCT attendance.emp_id)
        FROM attendance
        INNER JOIN users 
            ON CAST(attendance.emp_id AS INTEGER) = users.id
        WHERE attendance.date=?
    """, (today,)).fetchone()[0]

    total_count = len(corrected_rows)

    percentage = 0
    if total_users > 0:
        percentage = round((today_count / total_users) * 100, 2)

    report_default_start, report_default_end = get_default_report_dates()

    db.close()

    return render_template(
        "attendance.html",
        rows=corrected_rows,
        users=users,
        total_users=total_users,
        today_count=today_count,
        total_count=total_count,
        percentage=percentage,
        report_default_start=report_default_start,
        report_default_end=report_default_end,
    )

# ---------------- REDIRECTS FOR OLD ROUTES ----------------
@app.route("/attendance/students")
def student_attendance():
    return redirect("/attendance")

@app.route("/attendance/faculty")
def faculty_attendance():
    return redirect("/attendance")
@app.route("/attendance/analytics")
def attendance_analytics():

    db = get_db()

    # ---------------------------------------------------
    # GET ALL VALID SESSIONS (ONLY EXISTING USERS)
    # ---------------------------------------------------
    rows = db.execute("""
        SELECT users.name,
               users.role,
               attendance.emp_id,
               attendance.date,
               attendance.checkin_time,
               attendance.checkout_time,
               attendance.worked_hours
        FROM attendance
        INNER JOIN users ON attendance.emp_id = users.id
        WHERE attendance.checkout_time IS NOT NULL
        ORDER BY attendance.date DESC, attendance.checkin_time
    """).fetchall()

    # ---------------------------------------------------
    # MONTHLY SUMMARY
    # ---------------------------------------------------
    monthly_summary = db.execute("""
        SELECT strftime('%Y-%m', attendance.date) as month,
               ROUND(SUM(attendance.worked_hours),2)
        FROM attendance
        INNER JOIN users ON attendance.emp_id = users.id
        WHERE attendance.worked_hours IS NOT NULL
        GROUP BY month
        ORDER BY month DESC
    """).fetchall()

    # ---------------------------------------------------
    # HEATMAP DATA (Daily attendance count)
    # ---------------------------------------------------
    heatmap_data = db.execute("""
        SELECT attendance.date,
               COUNT(DISTINCT attendance.emp_id)
        FROM attendance
        INNER JOIN users ON attendance.emp_id = users.id
        GROUP BY attendance.date
        ORDER BY attendance.date DESC
    """).fetchall()

    # ---------------------------------------------------
    # LEADERBOARD (Top Employees)
    # ---------------------------------------------------
    leaderboard_raw = db.execute("""
        SELECT users.name,
               attendance.emp_id,
               ROUND(SUM(attendance.worked_hours),2) as total_hours
        FROM attendance
        INNER JOIN users ON attendance.emp_id = users.id
        WHERE attendance.worked_hours IS NOT NULL
        GROUP BY attendance.emp_id
        ORDER BY total_hours DESC
    """).fetchall()

    # Add ranking number
    leaderboard = []
    rank = 1
    for name, emp_id, total in leaderboard_raw:
        leaderboard.append({
            "rank": rank,
            "name": name,
            "emp_id": emp_id,
            "total_hours": total
        })
        rank += 1

    # ---------------------------------------------------
    # DAILY TOTAL + OVERTIME + PERFORMANCE
    # ---------------------------------------------------
    daily_data = {}

    for name, role, emp_id, date, checkin, checkout, worked in rows:

        if worked is None:
            continue

        key = (emp_id, date)

        if key not in daily_data:
            daily_data[key] = {
                "name": name,
                "role": role,
                "emp_id": emp_id,
                "date": date,
                "total_hours": 0,
                "sessions": []
            }

        daily_data[key]["total_hours"] += worked
        daily_data[key]["sessions"].append({
            "checkin": checkin,
            "checkout": checkout,
            "hours": round(worked, 2)
        })

    HOURLY_RATE = 200  # change if needed

    for key in daily_data:

        total = daily_data[key]["total_hours"]
        overtime = max(0, total - 8)
        overtime_salary = overtime * HOURLY_RATE * 1.5

        # AI Productivity Prediction Logic
        if total >= 9:
            score = "Excellent ⭐⭐⭐"
            prediction = "Highly Productive"
        elif total >= 8:
            score = "Good ⭐⭐"
            prediction = "Consistent Performer"
        elif total >= 6:
            score = "Average ⭐"
            prediction = "Moderate Productivity"
        else:
            score = "Low ⚠"
            prediction = "Needs Improvement"

        daily_data[key]["total_hours"] = round(total, 2)
        daily_data[key]["overtime"] = round(overtime, 2)
        daily_data[key]["overtime_salary"] = round(overtime_salary, 2)
        daily_data[key]["score"] = score
        daily_data[key]["prediction"] = prediction

    db.close()

    # Convert to sorted list (latest first)
    analytics_data = sorted(
        daily_data.values(),
        key=lambda x: x["date"],
        reverse=True
    )

    return render_template(
        "attendance_analytics.html",
        data=analytics_data,
        monthly_summary=monthly_summary,
        heatmap_data=heatmap_data,
        leaderboard=leaderboard
    )


@app.route("/download/pdf")
@app.route("/attendance/report/download")
def download_attendance_report():
    db = get_db()

    employee_id = parse_int_or_none(request.args.get("employee_id") or request.args.get("user_id"))
    start_raw = (request.args.get("start_date") or "").strip()
    end_raw = (request.args.get("end_date") or "").strip()
    default_start, default_end = get_default_report_dates()

    start_date_str = start_raw or default_start
    end_date_str = end_raw or default_end

    try:
        start_date = parse_iso_date(start_date_str)
        end_date = parse_iso_date(end_date_str)
    except ValueError:
        db.close()
        return "Invalid date range. Please select valid start and end dates.", 400

    if end_date < start_date:
        start_date, end_date = end_date, start_date

    if employee_id is not None:
        employee_exists = db.execute("SELECT id FROM users WHERE id = ?", (employee_id,)).fetchone()
        if not employee_exists:
            db.close()
            return "Selected employee was not found.", 404

    report_data = build_custom_attendance_report_data(db, employee_id, start_date, end_date)
    export_format = normalize_report_format(request.args.get("format"))

    if export_format == "xlsx":
        buffer = render_attendance_report_xlsx(report_data)
        mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        extension = "xlsx"
    elif export_format == "docx":
        buffer = render_attendance_report_docx(report_data)
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        extension = "docx"
    elif export_format == "png":
        buffer = render_attendance_report_png(report_data)
        mime_type = "image/png"
        extension = "png"
    else:
        buffer = render_attendance_report_pdf(report_data)
        mime_type = "application/pdf"
        extension = "pdf"

    db.close()

    return send_file(
        buffer,
        as_attachment=True,
        mimetype=mime_type,
        download_name=build_report_download_name(report_data["scope_label"], start_date, end_date, extension),
    )


@app.route("/payroll")
def payroll_dashboard():
    if not session.get("admin"):
        return redirect(url_for("login"))

    db = get_db()
    month_key = request.args.get("month") or get_month_key()

    available_months = [row[0] for row in db.execute("""
        SELECT DISTINCT substr(date, 1, 7) AS month
        FROM attendance
        WHERE date IS NOT NULL
        ORDER BY month DESC
    """).fetchall()]

    if month_key not in available_months:
        available_months.insert(0, month_key)

    payroll_rows, totals, month_label = build_payroll_snapshot(db, month_key)
    db.commit()
    db.close()

    msg = request.args.get("msg", "")
    return render_template(
        "payroll.html",
        payroll_rows=payroll_rows,
        totals=totals,
        month_key=month_key,
        month_label=month_label,
        available_months=available_months,
        msg=msg,
        current_month=get_month_key(),
    )


@app.route("/payroll/update_settings/<int:user_id>", methods=["POST"])
def update_payroll_settings(user_id):
    if not session.get("admin"):
        return redirect(url_for("login"))

    def to_float(field_name, default=0.0):
        value = request.form.get(field_name, "").strip()
        return float(value) if value else default

    def to_int(field_name, default=0):
        value = request.form.get(field_name, "").strip()
        return int(value) if value else default

    monthly_salary = to_float("monthly_salary")
    hourly_rate = to_float("hourly_rate")
    workdays_per_month = to_int("workdays_per_month")
    standard_hours = to_float("standard_hours", 8.0)
    overtime_multiplier = to_float("overtime_multiplier", 1.5)
    deduction_per_absent_day = to_float("deduction_per_absent_day")
    bonus_amount = to_float("bonus_amount")

    db = get_db()
    upsert_payroll_settings(
        db,
        user_id,
        monthly_salary,
        hourly_rate,
        workdays_per_month,
        standard_hours,
        overtime_multiplier,
        deduction_per_absent_day,
        bonus_amount,
    )
    db.commit()
    db.close()

    month_key = request.form.get("month") or get_month_key()
    redirect_to = request.form.get("redirect_to") or url_for("payroll_dashboard", month=month_key)
    separator = "&" if "?" in redirect_to else "?"
    return redirect(f"{redirect_to}{separator}msg=settings_saved")


@app.route("/api/payroll/summary")
def payroll_summary_api():
    if not session.get("admin"):
        return jsonify({"status": "fail", "error": "unauthorized"}), 403

    month_key = request.args.get("month") or get_month_key()
    db = get_db()
    payroll_rows, totals, month_label = build_payroll_snapshot(db, month_key)
    db.commit()
    db.close()

    return jsonify({
        "status": "ok",
        "month": month_key,
        "month_label": month_label,
        "totals": totals,
        "rows": payroll_rows,
    })


@app.route("/payroll/export/pdf")
def payroll_export_pdf():
    if not session.get("admin"):
        return redirect(url_for("login"))

    month_key = request.args.get("month") or get_month_key()
    db = get_db()
    payroll_rows, totals, month_label = build_payroll_snapshot(db, month_key)
    db.commit()
    db.close()

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5 * inch, bottomMargin=0.5 * inch)
    story = []
    styles = getSampleStyleSheet()

    story.append(Paragraph(f"<b>Payroll Statement - {month_label}</b>", styles["Title"]))
    story.append(Paragraph(f"<i>Generated on: {datetime.datetime.now().strftime('%B %d, %Y at %I:%M %p')}</i>", styles["Normal"]))
    story.append(Spacer(1, 0.2 * inch))

    summary_data = [["Employees", "Paid Days", "Present Days", "Absent Days", "Total Hours", "Overtime Hours", "Gross Salary", "Net Salary"]]
    summary_data.append([
        str(len(payroll_rows)),
        str(totals["paid_days"]),
        str(totals["present_days"]),
        str(totals["absent_days"]),
        f"{totals['total_hours']:.2f}",
        f"{totals['overtime_hours']:.2f}",
        f"{totals['gross_salary']:.2f}",
        f"{totals['net_salary']:.2f}",
    ])

    summary_table = Table(summary_data, colWidths=[0.8 * inch, 0.8 * inch, 0.85 * inch, 0.85 * inch, 0.85 * inch, 0.95 * inch, 1.0 * inch, 1.0 * inch])
    summary_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0d6efd")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("BACKGROUND", (0, 1), (-1, -1), colors.whitesmoke),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 0.25 * inch))

    detail_data = [["ID", "Name", "Role", "Paid", "Present", "Absent", "Hours", "Gross", "Deduction", "Net"]]
    for row in payroll_rows:
        detail_data.append([
            str(row["user_id"]),
            row["name"],
            row["role"],
            str(row["paid_days"]),
            str(row["present_days"]),
            str(row["absent_days"]),
            f"{row['total_hours']:.2f}",
            f"{row['gross_salary']:.2f}",
            f"{row['deductions']:.2f}",
            f"{row['net_salary']:.2f}",
        ])

    detail_table = Table(detail_data, colWidths=[0.45 * inch, 1.2 * inch, 0.7 * inch, 0.5 * inch, 0.55 * inch, 0.55 * inch, 0.75 * inch, 0.75 * inch, 0.75 * inch, 0.75 * inch])
    detail_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#198754")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 9),
        ("FONTSIZE", (0, 1), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
    ]))
    story.append(detail_table)

    doc.build(story)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"payroll_statement_{month_key}.pdf"
    )

@app.route("/enroll", methods=["GET", "POST"])
def enroll():
    if request.method == "POST":
        user_id = request.form["user_id"].strip()
        name = request.form["name"].strip()
        role = "employee"
        
        db = get_db()

        existing_user = db.execute("""
            SELECT id, name FROM users WHERE id = ?
        """, (user_id,)).fetchone()

        training_images = [f for f in os.listdir("TrainingImage") if f.startswith(f"User.{user_id}.")] if os.path.exists("TrainingImage") else []

        # Re-enrollment policy: allow existing user IDs to refresh their face samples.
        if existing_user:
            db.execute("""
                UPDATE users
                SET name = ?, role = ?
                WHERE id = ?
            """, (name, role, user_id))
            db.commit()
            db.close()
            return jsonify({
                "status": "ok",
                "message": f"User {user_id} updated from '{existing_user[1]}' to '{name}'. Ready for fresh face capture.",
                "existing_user_id": int(existing_user[0]),
                "existing_user_name": existing_user[1],
                "user_id": int(user_id),
                "reenroll": True,
            })

        # If user was deleted from DB but files remain, auto-clean orphan images and continue.
        if (not existing_user) and training_images:
            for img in training_images:
                try:
                    os.remove(os.path.join("TrainingImage", img))
                except Exception:
                    pass
        
        # Insert new user
        db.execute("""
            INSERT INTO users (id, name, role)
            VALUES (?, ?, ?)
        """, (user_id, name, role))
        db.commit()
        db.close()

        # Return JSON for AJAX workflow; client triggers /capture separately
        return jsonify({
            "status": "ok",
            "message": "User created. Ready for face capture.",
            "user_id": int(user_id)
        })

    return render_template("enroll.html")

@app.route("/check_user/<int:user_id>")
def check_user(user_id):
    """Check if user already exists"""
    db = get_db()
    existing_user = db.execute("""
        SELECT id, name FROM users WHERE id = ?
    """, (user_id,)).fetchone()
    db.close()
    
    # Check if training images exist
    training_images = [f for f in os.listdir("TrainingImage") if f.startswith(f"User.{user_id}.")] if os.path.exists("TrainingImage") else []
    
    if existing_user:
        return jsonify({
            "exists": True,
            "user_name": existing_user[1],
            "orphan_images": False,
            "allow_reenroll": True,
            "message": f"User ID {user_id} exists as '{existing_user[1]}'. Re-enrollment will replace old face samples.",
        })

    if training_images:
        # Orphan images (DB row deleted manually) should not block enrollment.
        return jsonify({"exists": False, "orphan_images": True})

    return jsonify({"exists": False, "orphan_images": False})

@app.route("/capture/<int:user_id>")
def capture_faces(user_id):
    global recognizer, camera

    # Release any existing camera connection
    if camera is not None:
        try:
            camera.release()
            camera = None
            time.sleep(0.5)  # Give camera time to release
        except:
            pass

    # Try to open camera with retries.
    # CAP_DSHOW only works on Windows; on Linux/macOS it can fail and prevent enrollment.
    cam = None
    if os.name == "nt":
        backend_candidates = [cv2.CAP_DSHOW, cv2.CAP_MSMF, cv2.CAP_ANY]
    else:
        backend_candidates = [cv2.CAP_ANY]

    for backend in backend_candidates:
        for attempt in range(3):
            cam = cv2.VideoCapture(0, backend)
            if cam.isOpened():
                break

            if cam is not None:
                cam.release()
            time.sleep(0.3)

        if cam is not None and cam.isOpened():
            break
    
    if cam is None or not cam.isOpened():
        print("❌ ERROR: Could not access camera")
        return jsonify({"status": "error", "trained": False, "message": "Camera access failed. Close other apps using camera."}), 500

    # Optimize camera settings for faster capture
    cam.set(cv2.CAP_PROP_FRAME_WIDTH, 640)
    cam.set(cv2.CAP_PROP_FRAME_HEIGHT, 480)
    cam.set(cv2.CAP_PROP_FPS, 30)
    cam.set(cv2.CAP_PROP_BUFFERSIZE, 1)
    
    # Warm up camera with a few reads
    for _ in range(5):
        cam.read()

    # Load face detector and verify it exists
    face_detector = cv2.CascadeClassifier("haarcascade_frontalface_default.xml")
    if face_detector.empty():
        cam.release()
        print("❌ ERROR: Haar Cascade file not found or invalid")
        return jsonify({"status": "error", "trained": False, "message": "Face detection model not found. Check haarcascade file."}), 500

    count = 0
    captured_faces = []
    detected_user_id = None
    detected_user_name = None
    duplicate_debug = None
    os.makedirs("TrainingImage", exist_ok=True)

    # Clear stale images for this user to avoid mixing old/new identity samples.
    clear_training_images_for_user(user_id)
    
    TARGET_IMAGES = 12  # Reduced from 20 for speed (still enough for training)
    max_attempts = 450  # Max ~15 seconds at 30fps (gives user enough time to position face)
    attempts = 0
    frame_skip = 0  # Capture every 2nd frame with face for variety
    no_face_count = 0  # Track frames without face detection

    print(f"[ENROLLMENT] Starting capture for user {user_id}...")
    start_time = time.time()

    while count < TARGET_IMAGES and attempts < max_attempts:
        attempts += 1
        
        ret, img = cam.read()
        if not ret:
            print(f"[ENROLLMENT] Warning: Failed to read frame {attempts}")
            continue

        gray_raw = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        gray = cv2.equalizeHist(gray_raw)
        # Reduced minNeighbors from 5 to 3 for easier face detection
        faces = face_detector.detectMultiScale(gray, scaleFactor=1.2, minNeighbors=3, minSize=(80, 80))
        if len(faces) == 0:
            # Fallback pass on raw grayscale for cameras where equalization hurts detection.
            faces = face_detector.detectMultiScale(gray_raw, scaleFactor=1.25, minNeighbors=4, minSize=(60, 60))
        if len(faces) > 1:
            faces = sorted(faces, key=lambda r: r[2] * r[3], reverse=True)

        if len(faces) > 1:
            if attempts % 30 == 0:
                print("[ENROLLMENT] Warning: Multiple faces detected. Please keep only one face in frame.")
            continue

        if len(faces) > 0:
            no_face_count = 0  # Reset no-face counter
            
            # Skip some frames for image variety (every 2nd detection)
            frame_skip += 1
            if frame_skip % 2 != 0:  # Capture every 2nd face detection
                continue
                
            for (x, y, w, h) in faces:
                if count >= TARGET_IMAGES:
                    break

                face_roi = gray[y:y+h, x:x+w]
                if face_roi.size == 0 or is_blurry(face_roi, threshold=85):
                    continue

                prepared_face = prepare_face_for_model(face_roi)
                if prepared_face is None:
                    continue

                count += 1
                captured_faces.append(prepared_face)

                # Save the CROPPED FACE ROI directly
                cv2.imwrite(
                    f"TrainingImage/User.{user_id}.{count}.jpg",
                    prepared_face
                )
                print(f"[ENROLLMENT] Captured image {count}/{TARGET_IMAGES}")

                # Early duplicate gate: run only after enough samples to avoid noisy false matches.
                if len(captured_faces) >= 8:
                    detected_user_id, detected_user_name, duplicate_debug = detect_duplicate_enrollment_face(captured_faces, user_id)
                    if detected_user_id:
                        print(f"[ENROLLMENT] Duplicate face detected early for user {user_id} -> existing ID {detected_user_id}")
                        attempts = max_attempts
                        break
                break  # Only capture one face per frame
        else:
            no_face_count += 1
            # Log warning if no face detected for extended period
            if no_face_count % 30 == 0:  # Every 30 frames (~1 second)
                print(f"[ENROLLMENT] Warning: No face detected for {no_face_count} frames. Please face the camera.")

    cam.release()
    elapsed_time = time.time() - start_time
    print(f"[ENROLLMENT] Finished: Captured {count}/{TARGET_IMAGES} images in {elapsed_time:.1f}s over {attempts} frames")

    # If duplicate face was detected during capture, block enrollment immediately.
    if detected_user_id:
        clear_training_images_for_user(user_id)

        return jsonify({
            "status": "duplicate_face",
            "message": f"⚠️ Duplicate ID Found: This face is already enrolled with ID {detected_user_id} ({detected_user_name}). Enrollment blocked.",
            "existing_user_id": detected_user_id,
            "existing_user_name": detected_user_name,
            "debug": duplicate_debug
        }), 400
    
    # Check if we got enough images
    if count < MIN_IMAGES_PER_USER:
        # Cleanup partial enrollment so user can retry immediately without manual DB edits.
        try:
            if os.path.exists("TrainingImage"):
                for img_file in os.listdir("TrainingImage"):
                    if img_file.startswith(f"User.{user_id}."):
                        os.remove(os.path.join("TrainingImage", img_file))
        except Exception:
            pass

        error_msg = f"Only captured {count}/{TARGET_IMAGES} images. Need at least {MIN_IMAGES_PER_USER}."
        
        if count == 0:
            error_msg += " No face detected. Ensure good lighting, remove glasses/mask, and face the camera directly."
        elif count < 4:
            error_msg += " Very few faces detected. Try better lighting and face the camera straight on."
        else:
            error_msg += " Please ensure good lighting and stay centered in the frame."
            
        print(f"[ENROLLMENT] ❌ {error_msg}")
        return jsonify({
            "status": "error", 
            "trained": False, 
            "message": error_msg
        }), 400

    # Check if captured face matches any existing user (detect fake IDs)
    detected_user_id = None
    detected_user_name = None
    duplicate_debug = None
    try:
        detected_user_id, detected_user_name, duplicate_debug = detect_duplicate_enrollment_face(captured_faces, user_id)
    except Exception as e:
        print(f"Face matching error: {e}")

    # If a different user's face was detected, return error with warning
    if detected_user_id:
        # Delete the wrongly saved images for this fake ID attempt
        clear_training_images_for_user(user_id)
        
        return jsonify({
            "status": "duplicate_face", 
            "message": f"⚠️ Duplicate ID Found: This face is already enrolled with ID {detected_user_id} ({detected_user_name}). Enrollment blocked.",
            "existing_user_id": detected_user_id,
            "existing_user_name": detected_user_name,
            "debug": duplicate_debug
        }), 400

    # Automatically train the model after capturing images
    print(f"[TRAINING] Starting training with {count} images...")
    training_start = time.time()
    
    if count == 0:
        return jsonify({"status": "captured", "trained": False, "message": "No faces detected. Please try again."})
    
    success = train_model()
    
    if success:
        total_time = time.time() - start_time
        print(f"✅ Total enrollment time: {total_time:.1f}s (capture: {elapsed_time:.1f}s, training: {total_time - elapsed_time:.1f}s)")
        
        # Reload the recognizer with the newly trained model
        try:
            reload_recognizer_from_disk()
            print("✓ Recognizer reloaded")
        except Exception as e:
            print(f"Warning: Could not reload recognizer: {e}")

        # Refresh caches so recognition immediately sees newly enrolled users.
        get_user_cache(force_refresh=True)
        get_training_image_counts(force_refresh=True)

        return jsonify({
            "status": "captured", 
            "trained": True, 
            "message": f"✅ Enrollment complete! ({count} images, {total_time:.1f}s)"
        })
    else:
        return jsonify({"status": "captured", "trained": False, "message": "Images captured but training failed. Please try again."})

@app.route("/set_threshold", methods=["POST"])
def set_threshold():
    global THRESHOLD
    if not admin_required():
        return "Unauthorized", 403

    requested = int(request.form["threshold"])
    THRESHOLD = max(35, min(70, requested))
    return redirect("/admin")


@app.route("/set_mask_config", methods=["POST"])
def set_mask_config():
    global MASK_AWARE_ENABLED
    global MASK_THRESHOLD_BOOST
    global MASK_REQUIRED_STABLE_MATCHES
    global MASK_MIN_UPPER_CORR
    global MASK_MIN_UPPER_GAP
    global MASK_MAX_UPPER_LBP_DIST
    global MASK_MIN_UPPER_LBP_MARGIN
    global IRIS_VERIFY_ENABLED
    global IRIS_MIN_SCORE
    global IRIS_MIN_GAP
    global IRIS_MIN_VOTES

    if not admin_required():
        return "Unauthorized", 403

    try:
        MASK_AWARE_ENABLED = request.form.get("mask_enabled") == "on"
        MASK_THRESHOLD_BOOST = max(0, min(20, int(request.form.get("mask_threshold_boost", MASK_THRESHOLD_BOOST))))
        MASK_REQUIRED_STABLE_MATCHES = max(2, min(8, int(request.form.get("mask_stable_matches", MASK_REQUIRED_STABLE_MATCHES))))
        MASK_MIN_UPPER_CORR = max(0.20, min(0.95, float(request.form.get("mask_min_upper_corr", MASK_MIN_UPPER_CORR))))
        MASK_MIN_UPPER_GAP = max(0.0, min(0.20, float(request.form.get("mask_min_upper_gap", MASK_MIN_UPPER_GAP))))
        MASK_MAX_UPPER_LBP_DIST = max(0.20, min(3.50, float(request.form.get("mask_max_upper_lbp", MASK_MAX_UPPER_LBP_DIST))))
        MASK_MIN_UPPER_LBP_MARGIN = max(0.0, min(0.50, float(request.form.get("mask_min_upper_lbp_margin", MASK_MIN_UPPER_LBP_MARGIN))))
        IRIS_VERIFY_ENABLED = request.form.get("iris_enabled") == "on"
        IRIS_MIN_SCORE = max(-0.20, min(1.00, float(request.form.get("iris_min_score", IRIS_MIN_SCORE))))
        IRIS_MIN_GAP = max(0.0, min(0.30, float(request.form.get("iris_min_gap", IRIS_MIN_GAP))))
        IRIS_MIN_VOTES = max(1, min(2, int(request.form.get("iris_min_votes", IRIS_MIN_VOTES))))
    except ValueError:
        return redirect("/admin?msg=mask_invalid")

    return redirect("/admin?msg=mask_saved")

@app.route("/stats")
def stats():
    db = get_db()
    data = db.execute("""
    SELECT COUNT(DISTINCT emp_id)
    GROUP BY date
    """).fetchall()
    db.close()
    return jsonify(data)

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")

        # 🔐 ADMIN CREDENTIALS (ACADEMIC DEMO)
        if username == "admin" and password == "admin123":
            session["admin"] = True
            return redirect(url_for("admin"))

        return render_template("login.html", error="Invalid credentials")

    return render_template("login.html")


@app.route("/admin")
def admin():
    if not session.get("admin"):
        return redirect(url_for("login"))   # 🔒 FORCE LOGIN

    db = get_db()
    total_users = db.execute("SELECT COUNT(*) FROM users").fetchone()[0]
    total_attendance = db.execute("SELECT COUNT(*) FROM attendance").fetchone()[0]
    pending_leave_requests = db.execute("""
        SELECT COUNT(*) FROM leave_requests WHERE request_status = 'Pending'
    """).fetchone()[0]
    recent_leave_requests = db.execute("""
        SELECT lr.id,
               lr.emp_id,
               COALESCE(u.name, 'Unknown') AS emp_name,
               COALESCE(u.role, 'employee') AS emp_role,
               lr.subject,
               lr.leave_type,
               lr.start_date,
               lr.end_date,
               lr.total_days,
               lr.reason,
               lr.request_status,
               COALESCE(lr.admin_comment, '') AS admin_comment,
               lr.applied_at,
               lr.reviewed_at,
               COALESCE(lr.reviewed_by, '-')
        FROM leave_requests lr
        LEFT JOIN users u ON CAST(lr.emp_id AS INTEGER) = u.id
        ORDER BY CASE lr.request_status WHEN 'Pending' THEN 0 WHEN 'Approved' THEN 1 ELSE 2 END,
                 lr.id DESC
        LIMIT 10
    """).fetchall()
    payroll_settings_rows = build_payroll_settings_rows(db, get_month_key())
    db.close()

    return render_template(
        "admin.html",
        total_users=total_users,
        total_attendance=total_attendance,
        pending_leave_requests=pending_leave_requests,
        recent_leave_requests=recent_leave_requests,
        payroll_settings_rows=payroll_settings_rows,
        current_month=get_month_key(),
        threshold=THRESHOLD,
        mask_aware_enabled=MASK_AWARE_ENABLED,
        mask_threshold_boost=MASK_THRESHOLD_BOOST,
        mask_required_stable_matches=MASK_REQUIRED_STABLE_MATCHES,
        mask_min_upper_corr=MASK_MIN_UPPER_CORR,
        mask_min_upper_gap=MASK_MIN_UPPER_GAP,
        mask_max_upper_lbp_dist=MASK_MAX_UPPER_LBP_DIST,
        mask_min_upper_lbp_margin=MASK_MIN_UPPER_LBP_MARGIN,
        iris_verify_enabled=IRIS_VERIFY_ENABLED,
        iris_min_score=IRIS_MIN_SCORE,
        iris_min_gap=IRIS_MIN_GAP,
        iris_min_votes=IRIS_MIN_VOTES
    )

@app.route("/admin/complaints")
def admin_complaints():
    if not session.get("admin"):
        return redirect(url_for("login"))

    db = get_db()
    complaints = db.execute("""
        SELECT id, category, description, status, date
        FROM complaints
        ORDER BY date DESC
    """).fetchall()
    db.close()

    return render_template("admin_complaints.html", complaints=complaints)


@app.route("/admin/close_complaint/<int:cid>")
def close_complaint(cid):
    if not admin_required():
        return redirect("/login")

    db = get_db()
    db.execute("UPDATE complaints SET status='Closed' WHERE id=?", (cid,))
    db.commit()
    db.close()

    return redirect("/admin/complaints")

@app.route("/admin/system_status")
def system_status():
    if not admin_required():
        return redirect("/login")

    avg_conf = 0
    if confidence_scores:
        avg_conf = round(sum(confidence_scores) / len(confidence_scores), 2)

    return jsonify({
        "camera_status": camera_status,
        "fps": fps_value,
        "success": success_recognition,
        "failure": failed_recognition,
        "avg_confidence": avg_conf
    })

@app.route("/stats/daily")
def daily_stats():
    db = get_db()
    data = db.execute("""
        SELECT 
            date,
            SUM(
                CASE
                    WHEN LOWER(COALESCE(status, '')) = 'leave approved' THEN -1
                    WHEN LOWER(COALESCE(status, 'present')) = 'present' THEN 1
                    ELSE 0
                END
            ) AS trend_score
        FROM attendance
        GROUP BY date
        ORDER BY date
    """).fetchall()
    db.close()
    return jsonify(data)


@app.route("/stats/performance")
def performance_stats():
    return jsonify({
        "success": success_recognition,
        "failure": failed_recognition
    })

@app.route("/feedback", methods=["GET", "POST"])
def feedback():
    if request.method == "POST":
        rating = request.form["rating"]
        message = request.form["message"]

        db = get_db()
        db.execute("""
            INSERT INTO feedback (rating, message, date)
            VALUES (?, ?, date('now'))
        """, (rating, message))
        db.commit()
        db.close()

    return render_template("feedback.html")

@app.route("/admin/feedback")
def admin_feedback():
    if not session.get("admin"):
        return redirect(url_for("login"))

    db = get_db()
    feedback = db.execute("""
        SELECT rating, message, date
        FROM feedback
        ORDER BY date DESC
    """).fetchall()
    db.close()

    return render_template("admin_feedback.html", feedback=feedback)


@app.route("/leave/request", methods=["GET", "POST"])
def leave_request():
    return redirect(url_for("support", section="leave"))


@app.route("/admin/leave_requests")
def admin_leave_requests():
    if not admin_required():
        return redirect("/login")

    msg = request.args.get("msg")
    applied = request.args.get("applied")
    skipped = request.args.get("skipped")

    db = get_db()
    requests = db.execute("""
        SELECT lr.id,
               lr.emp_id,
               COALESCE(u.name, 'Unknown') AS emp_name,
               COALESCE(u.role, 'employee') AS emp_role,
               lr.subject,
               lr.leave_type,
               lr.start_date,
               lr.end_date,
               lr.total_days,
               lr.reason,
               lr.request_status,
               COALESCE(lr.admin_comment, '') AS admin_comment,
               lr.applied_at,
               lr.reviewed_at,
               COALESCE(lr.reviewed_by, '-')
        FROM leave_requests lr
        LEFT JOIN users u ON CAST(lr.emp_id AS INTEGER) = u.id
        ORDER BY CASE lr.request_status WHEN 'Pending' THEN 0 WHEN 'Approved' THEN 1 ELSE 2 END,
                 lr.id DESC
    """).fetchall()
    db.close()

    return render_template(
        "admin_leave_requests.html",
        requests=requests,
        msg=msg,
        applied=applied,
        skipped=skipped,
    )


@app.route("/admin/leave_requests/<int:request_id>/action", methods=["POST"])
def admin_leave_request_action(request_id):
    if not admin_required():
        return redirect("/login")

    action = (request.form.get("action") or "").strip().lower()
    admin_comment = (request.form.get("admin_comment") or "").strip()
    if action not in {"approve", "deny"}:
        return redirect("/admin/leave_requests?msg=invalid_action")

    db = get_db()
    try:
        request_row = db.execute("""
            SELECT id, emp_id, start_date, end_date, request_status
            FROM leave_requests
            WHERE id = ?
        """, (request_id,)).fetchone()

        if not request_row:
            db.close()
            return redirect("/admin/leave_requests?msg=not_found")

        _, emp_id, start_date_str, end_date_str, current_status = request_row
        if current_status != "Pending":
            db.close()
            return redirect("/admin/leave_requests?msg=already_processed")

        reviewed_at = datetime.datetime.now().isoformat(timespec="seconds")

        if action == "approve":
            start_date = parse_iso_date(start_date_str)
            end_date = parse_iso_date(end_date_str)
            applied, skipped = apply_approved_leave_to_attendance(db, emp_id, start_date, end_date)
            db.execute("""
                UPDATE leave_requests
                SET request_status = 'Approved',
                    admin_comment = ?,
                    reviewed_at = ?,
                    reviewed_by = 'admin'
                WHERE id = ?
            """, (admin_comment, reviewed_at, request_id))
            db.commit()
            db.close()
            return redirect(f"/admin/leave_requests?msg=approved&applied={applied}&skipped={skipped}")

        db.execute("""
            UPDATE leave_requests
            SET request_status = 'Denied',
                admin_comment = ?,
                reviewed_at = ?,
                reviewed_by = 'admin'
            WHERE id = ?
        """, (admin_comment, reviewed_at, request_id))
        db.commit()
        db.close()
        return redirect("/admin/leave_requests?msg=denied")
    except Exception:
        db.close()
        return redirect("/admin/leave_requests?msg=error")



@app.route("/grievance", methods=["GET", "POST"])
def grievance():
    if request.method == "POST":
        category = request.form["category"]
        description = request.form["description"]

        db = get_db()
        db.execute("""
            INSERT INTO complaints (category, description, date)
            VALUES (?, ?, date('now'))
        """, (category, description))
        db.commit()
        db.close()

    return render_template("grievance.html")

@app.route("/faq")
def faq():
    return render_template("faq.html")

@app.route("/camera/on")
def camera_on():
    global camera_enabled
    camera_enabled = True
    return jsonify({"status": "ON"})

@app.route("/camera/off")
def camera_off():
    global camera_enabled, camera
    camera_enabled = False

    if camera is not None:
        camera.release()
        camera = None

    return jsonify({"status": "OFF"})

@app.route("/test_preview")
def test_preview():
    cam = cv2.VideoCapture(0)
    success, frame = cam.read()
    cam.release()

    if not success:
        return jsonify({"status": "fail"})

    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
    faces = face_cascade.detectMultiScale(gray, 1.3, 5)

    result = "No face detected"

    for (x,y,w,h) in faces:
        # Reload latest model before recognition
        if os.path.exists("TrainingImageLabel/Trainner.yml"):
            recognizer.read("TrainingImageLabel/Trainner.yml")
        # Use the same preprocessing as enrollment and live recognition.
        face_roi_resized = prepare_face_for_model(gray[y:y+h, x:x+w])
        if face_roi_resized is None:
            continue
        uid, conf, prediction_debug = predict_face_with_flip(recognizer, face_roi_resized)
        if uid is None or conf is None:
            continue
        gallery_ok, _ = verify_prediction_with_gallery(face_roi_resized, uid)
        if conf <= min(THRESHOLD, 52) and prediction_debug.get("agreement", False) and gallery_ok:
            result = f"Recognized: User {uid}"
        else:
            result = "Face detected but not recognized"

    return jsonify({"status": "ok", "result": result})

@app.route("/admin/model_accuracy")
def model_accuracy():
    total = success_recognition + failed_recognition
    accuracy = round((success_recognition / total) * 100, 2) if total else 0

    return jsonify({
        "accuracy": accuracy,
        "success": success_recognition,
        "failure": failed_recognition
    })


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

@app.route("/admin/accuracy_data")
def accuracy_data():

    date_filter = request.args.get("date")

    db = get_db()
    cur = db.cursor()

    if date_filter:
        cur.execute("""
            SELECT strftime('%H:00', date || ' ' || checkin_time) as hour,
                   SUM(CASE WHEN status != 'failed' THEN 1 ELSE 0 END) as success,
                   SUM(CASE WHEN status = 'failed' THEN 1 ELSE 0 END) as failure
            FROM attendance
            WHERE date = ?
            GROUP BY hour
            ORDER BY hour
        """, (date_filter,))
    else:
        cur.execute("""
            SELECT strftime('%H:00', date || ' ' || checkin_time) as hour,
                   SUM(CASE WHEN status != 'failed' THEN 1 ELSE 0 END) as success,
                   SUM(CASE WHEN status = 'failed' THEN 1 ELSE 0 END) as failure
            FROM attendance
            GROUP BY hour
            ORDER BY hour
        """)

    rows = cur.fetchall()
    db.close()

    result = []

    for hour, success, failure in rows:
        total = (success or 0) + (failure or 0)

        accuracy = round((success / total) * 100, 2) if total else 0

        result.append({
            "hour": hour,
            "accuracy": accuracy
        })

    return jsonify(result)
@app.route("/admin/success_failure_data")
def success_failure_data():
    db = get_db()
    data = db.execute("""
        SELECT date,
               SUM(CASE WHEN status!='failed' THEN 1 ELSE 0 END),
               SUM(CASE WHEN status='failed' THEN 1 ELSE 0 END)
        FROM attendance
        GROUP BY date
    """).fetchall()
    db.close()

    return jsonify({
        "labels": [d[0] for d in data],
        "success": [d[1] for d in data],
        "failure": [d[2] for d in data]
    })
@app.route("/admin/confidence_distribution")
def confidence_distribution():

    # Example ranges
    labels = ["0-50", "50-70", "70-85", "85-100"]
    values = [5, 10, 25, 40]  # Replace with real stored confidence

    return jsonify({
        "labels": labels,
        "values": values
    })



@app.route("/admin/manual_attendance", methods=["POST"])
def manual_attendance():
    emp_id = request.form.get("emp_id")
    date = request.form.get("date")
    checkin = request.form.get("checkin")

    db = get_db()
    db.execute("""
        INSERT INTO attendance (emp_id, date, checkin_time)
        VALUES (?, ?, ?)
    """, (emp_id, date, checkin))

    db.commit()
    db.close()
    return redirect("/attendance")

@app.route("/attendance_status")
def attendance_status_api():
    global attendance_status
    status = attendance_status
    attendance_status = ""  # reset after reading
    return jsonify({"status": status})

@app.route("/support", methods=["GET", "POST"])
def support():

    db = get_db()
    msg = request.args.get("msg")
    section = request.args.get("section", "support")
    leave_msg = request.args.get("leave_msg")
    leave_error = None

    if request.method == "POST":

        form_type = request.form.get("type")
        message = request.form.get("message")
        rating = request.form.get("rating")
        category = request.form.get("category")

        if form_type == "feedback":
            db.execute("""
                INSERT INTO feedback (rating, message, date)
                VALUES (?, ?, date('now'))
            """, (rating, message))

        elif form_type == "grievance":
            db.execute("""
                INSERT INTO complaints (category, description, date)
                VALUES (?, ?, date('now'))
            """, (category, message))

        elif form_type == "leave_request":
            try:
                submit_leave_request(
                    db,
                    request.form.get("emp_id"),
                    request.form.get("subject"),
                    request.form.get("leave_type"),
                    request.form.get("start_date"),
                    request.form.get("end_date"),
                    request.form.get("reason"),
                )
                db.commit()
                db.close()
                return redirect("/support?leave_msg=submitted&section=leave")
            except ValueError as ve:
                leave_error = str(ve)
            except Exception:
                leave_error = "Could not submit leave request. Please try again."

        db.commit()

    users = db.execute("""
        SELECT id, name, role
        FROM users
        ORDER BY name
    """).fetchall()

    leave_requests = db.execute("""
        SELECT lr.id,
               lr.emp_id,
               COALESCE(u.name, 'Unknown') AS emp_name,
               lr.subject,
               lr.leave_type,
               lr.start_date,
               lr.end_date,
               lr.total_days,
               lr.request_status,
               lr.applied_at,
               lr.reviewed_at,
               COALESCE(lr.admin_comment, '')
        FROM leave_requests lr
        LEFT JOIN users u ON CAST(lr.emp_id AS INTEGER) = u.id
        ORDER BY lr.id DESC
        LIMIT 50
    """).fetchall()
    db.close()

    return render_template(
        "support.html",
        msg=msg,
        section=section,
        leave_msg=leave_msg,
        leave_error=leave_error,
        users=users,
        leave_requests=leave_requests,
    )

@app.route("/set_mode/<mode>")
def set_mode(mode):
    global attendance_mode

    if mode in ["checkin", "checkout"]:
        attendance_mode = mode

    return jsonify({"mode": attendance_mode})


@app.route("/copilot")
def copilot():
    return render_template("copilot.html")
@app.route("/ai_assistant", methods=["POST"])
def ai_assistant():

    user_msg = request.json.get("message").lower()

    if "check in" in user_msg:
        reply = "To check in, open Live → Check-In Session and complete face verification."

    elif "check out" in user_msg:
        reply = "To check out, open Live → Check-Out Session and complete verification."

    elif "attendance percentage" in user_msg:
        reply = "Attendance percentage = (Today's Attendance / Total Users) × 100."

    elif "camera not working" in user_msg:
        reply = "Make sure no other application is using the camera and refresh the page."

    elif "delete record" in user_msg:
        reply = "Go to Attendance page and use the delete option next to the record."

    elif "payroll" in user_msg or "salary" in user_msg:
        reply = "Open the Payroll page to configure salary settings and generate monthly pay from attendance data."

    else:
        reply = "I'm FAS Copilot 🤖. I can help with attendance, admin tools, analytics, reports, and troubleshooting."

    return jsonify({"reply": reply})
@app.route("/api/copilot", methods=["POST"])
def copilot_api():
    data = request.json
    message = data.get("message", "").lower()

    reply = "I'm FAS Copilot 🤖. How can I help you?"

    # Attendance help
    if "attendance" in message:
        reply = "You can view attendance in the Attendance page. Use filters for daily, weekly, or monthly insights."

    elif "check in" in message:
        reply = "To check-in, select Check-In Session under Live and start the camera."

    elif "check out" in message:
        reply = "To check-out, select Check-Out Session under Live and verify your face."

    elif "admin" in message:
        reply = "Admin dashboard provides analytics, reports, complaints management, and accuracy monitoring."

    elif "accuracy" in message:
        reply = "Model accuracy is calculated using successful recognitions divided by total recognitions."

    elif "report" in message:
        reply = "You can download PDF reports from Attendance or Admin dashboard."

    elif "payroll" in message or "salary" in message:
        reply = "Payroll is available from the Payroll dashboard. It calculates monthly salary, absences, and overtime from attendance."

    elif "support" in message or "complaint" in message:
        reply = "Use Support Center to submit complaints or feedback."

    elif "user" in message:
        reply = "You can enroll new users from the Enroll page."

    return jsonify({"reply": reply})

@app.route("/admin/users")
def admin_users():

    if not session.get("admin"):
        return redirect(url_for("login"))

    db = get_db()
    users = db.execute("""
        SELECT id, name, role
        FROM users
        ORDER BY id
    """).fetchall()
    db.close()

    msg = request.args.get("msg", "")
    return render_template("admin_users.html", users=users, msg=msg)
@app.route("/admin/edit_user/<int:user_id>", methods=["POST"])
def edit_user(user_id):

    if not session.get("admin"):
        return redirect(url_for("login"))

    name = request.form.get("name")
    role = (request.form.get("role") or "employee").strip().lower()
    if role not in ("admin", "employee"):
        role = "employee"

    db = get_db()
    db.execute("""
        UPDATE users
        SET name=?, role=?
        WHERE id=?
    """, (name, role, user_id))
    db.commit()
    db.close()

    return redirect("/admin/users")
@app.route("/admin/delete_user/<int:user_id>", methods=["POST"])
def delete_user(user_id):

    if not session.get("admin"):
        return redirect(url_for("login"))

    db = get_db()

    # Prevent deleting main admin
    if user_id == 1:
        db.close()
        return redirect("/admin/users?msg=protected")

    # Delete attendance
    db.execute("DELETE FROM attendance WHERE emp_id=?", (user_id,))

    # Delete user
    cur = db.execute("DELETE FROM users WHERE id=?", (user_id,))
    deleted_rows = cur.rowcount

    db.commit()
    db.close()

    # Delete training images
    import os
    folder = "TrainingImage"

    if os.path.exists(folder):
        for file in os.listdir(folder):
            if file.startswith(f"User.{user_id}."):
                os.remove(os.path.join(folder, file))

    # Retrain model safely
    if os.path.exists("TrainingImageLabel/Trainner.yml"):
        try:
            train_model()
        except Exception:
            pass

    if deleted_rows and deleted_rows > 0:
        return redirect("/admin/users?msg=deleted")

    return redirect("/admin/users?msg=not_found")

# =============== DIAGNOSTIC ENDPOINTS ===============
@app.route("/debug/training_status")
def training_status():
    """Check current training status and model with detailed diagnostics"""
    try:
        status = {
            "training_images": {},
            "database_users": [],
            "model_exists": os.path.exists("TrainingImageLabel/Trainner.yml"),
            "diagnostics": [],
            "model_info": {}
        }
        
        # Check training images
        if os.path.exists("TrainingImage"):
            images = os.listdir("TrainingImage")
            for img in images:
                if img.endswith(".jpg"):
                    parts = img.split(".")
                    if len(parts) >= 2:
                        user_id = parts[1]
                        if user_id not in status["training_images"]:
                            status["training_images"][user_id] = []
                        status["training_images"][user_id].append(img)
            
            # Convert to count
            training_counts = {uid: len(images) for uid, images in status["training_images"].items()}
            status["training_images"] = training_counts
        
        # Check database users
        db = get_db()
        users = db.execute("SELECT id, name, role FROM users ORDER BY id").fetchall()
        status["database_users"] = [{"id": u[0], "name": u[1], "role": u[2]} for u in users]
        db.close()
        
        # Validation checks
        if not status["training_images"]:
            status["diagnostics"].append("❌ No training images found!")
        
        for user_id_str, count in status["training_images"].items():
            user_id = int(user_id_str)
            user_exists = any(u["id"] == user_id for u in status["database_users"])
            
            if not user_exists:
                status["diagnostics"].append(f"⚠️  Training images exist for user {user_id} but NOT in database!")
            elif count < 10:
                status["diagnostics"].append(f"⚠️  User {user_id} has only {count} training images (need 10+)")
            else:
                status["diagnostics"].append(f"✅ User {user_id} has {count} training images")
        
        for user in status["database_users"]:
            user_id_str = str(user["id"])
            if user_id_str not in status["training_images"]:
                status["diagnostics"].append(f"❌ User {user['id']} ({user['name']}) in DB but NO training images!")
        
        # Check model info if it exists
        if status["model_exists"]:
            try:
                test_recognizer = cv2.face.LBPHFaceRecognizer_create()
                test_recognizer.read("TrainingImageLabel/Trainner.yml")
                status["model_info"] = {
                    "status": "✅ Model file can be loaded",
                    "file_size": os.path.getsize("TrainingImageLabel/Trainner.yml"),
                    "last_modified": os.path.getmtime("TrainingImageLabel/Trainner.yml")
                }
            except Exception as e:
                status["model_info"] = {
                    "status": f"❌ Error loading model: {str(e)}"
                }
        
        return jsonify(status)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/debug/test_model", methods=["POST"])
def test_model():
    """Test the trained model against actual training images"""
    try:
        if not os.path.exists("TrainingImageLabel/Trainner.yml"):
            return jsonify({"status": "error", "message": "No model file found. Train the model first."}), 400
        
        # Load the recognizer
        test_recognizer = cv2.face.LBPHFaceRecognizer_create()
        test_recognizer.read("TrainingImageLabel/Trainner.yml")
        
        results = {
            "status": "ok",
            "tests": [],
            "accuracy": 0
        }
        
        if not os.path.exists("TrainingImage"):
            return jsonify({"status": "error", "message": "No training images found"}), 400
        
        all_images = [f for f in os.listdir("TrainingImage") if f.endswith(".jpg")]
        correct_predictions = 0
        
        # Test first 3 images from each user
        user_images = {}
        for img in all_images:
            parts = img.split(".")
            if len(parts) >= 2:
                user_id_str = parts[1]
                if user_id_str not in user_images:
                    user_images[user_id_str] = []
                user_images[user_id_str].append(img)
        
        for user_id_str in sorted(user_images.keys()):
            actual_user_id = int(user_id_str)
            images = user_images[user_id_str][:3]  # Test first 3 images
            
            for img_name in images:
                try:
                    # Load and prepare image
                    img_path = os.path.join("TrainingImage", img_name)
                    gray_img = Image.open(img_path).convert("L")
                    img_np = np.array(gray_img, "uint8")
                    img_resized = cv2.resize(img_np, (200, 200))
                    
                    # Predict
                    predicted_id, conf = test_recognizer.predict(img_resized)
                    
                    is_correct = (predicted_id == actual_user_id)
                    if is_correct:
                        correct_predictions += 1
                    
                    results["tests"].append({
                        "image": img_name,
                        "actual_user_id": actual_user_id,
                        "predicted_user_id": predicted_id,
                        "confidence": round(conf, 2),
                        "correct": is_correct,
                        "status": "✅ MATCH" if is_correct else "❌ MISMATCH"
                    })
                except Exception as e:
                    results["tests"].append({
                        "image": img_name,
                        "error": str(e)
                    })
        
        if results["tests"]:
            results["accuracy"] = round((correct_predictions / len(results["tests"])) * 100, 2)
        
        return jsonify(results)
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route("/debug/retrain", methods=["POST"])
def retrain_model():
    """Force retrain the model with maximum diagnostics"""
    try:
        # Delete old model to ensure clean training
        if os.path.exists("TrainingImageLabel/Trainner.yml"):
            os.remove("TrainingImageLabel/Trainner.yml")
            print("[RETRAIN] Deleted old model file")
        
        print("[RETRAIN] Starting model training from scratch...")
        result = train_model()
        
        if not result:
            return jsonify({
                "status": "error", 
                "message": "❌ Training failed - check Flask console for details"
            }), 400
        
        # Verify model was created
        if not os.path.exists("TrainingImageLabel/Trainner.yml"):
            return jsonify({
                "status": "error",
                "message": "❌ Model file not created after training"
            }), 400
        
        # Reload the global recognizer
        global recognizer
        recognizer = cv2.face.LBPHFaceRecognizer_create()
        recognizer.read("TrainingImageLabel/Trainner.yml")
        print("[RETRAIN] ✅ Global recognizer reloaded with new model")
        
        return jsonify({
            "status": "success",
            "message": "✅ Model retrained and reloaded successfully. Test with /debug/test_model"
        })
    except Exception as e:
        print(f"[ERROR] Retrain failed: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route("/debug/attendance_records")
def debug_attendance_records():
    """Check all attendance records in database"""
    try:
        db = get_db()
        
        # Get raw attendance records
        records = db.execute("""
            SELECT id, emp_id, date, day, checkin_time, checkout_time, worked_hours
            FROM attendance
            ORDER BY id DESC
            LIMIT 20
        """).fetchall()
        
        # Get user info
        users = db.execute("SELECT id, name FROM users").fetchall()
        db.close()
        
        user_map = {str(u[0]): u[1] for u in users}
        
        result = {
            "total_records": len(records),
            "recent_records": [],
            "user_map": user_map
        }
        
        for rec in records:
            emp_id_db = rec[1]  # What's stored in database
            
            # Try to get user name from database
            db = get_db()
            user_name_query = db.execute(
                "SELECT name FROM users WHERE id = ?",
                (emp_id_db,)
            ).fetchone()
            db.close()
            
            user_name_db = user_name_query[0] if user_name_query else "NOT FOUND"
            
            result["recent_records"].append({
                "attendance_id": rec[0],
                "emp_id_stored": emp_id_db,
                "user_name_from_db": user_name_db,
                "date": rec[2],
                "checkin_time": rec[4],
                "checkout_time": rec[5],
                "worked_hours": rec[6]
            })
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500
def clear_all_training():
    """DANGEROUS: Clear ALL training data and start fresh"""
    try:
        import shutil
        
        # Delete training images folder
        if os.path.exists("TrainingImage"):
            shutil.rmtree("TrainingImage")
            print("[CLEANUP] Deleted TrainingImage folder")
        
        # Delete model folder
        if os.path.exists("TrainingImageLabel"):
            shutil.rmtree("TrainingImageLabel")
            print("[CLEANUP] Deleted TrainingImageLabel folder")
        
        # Delete all users from database
        db = get_db()
        db.execute("DELETE FROM users")
        db.execute("DELETE FROM attendance")
        db.commit()
        db.close()
        print("[CLEANUP] Cleared all users and attendance records from database")
        
        return jsonify({
            "status": "success",
            "message": "✅ All training data, users, and attendance records have been cleared. You can now start fresh with enrollment."
        })
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

# ------------------ Main ------------------
if __name__ == "__main__":
    app.run(debug=True)

